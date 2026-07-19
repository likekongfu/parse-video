import asyncio
import io
import inspect
import json
import logging
from collections import Counter
from pathlib import Path
from unittest.mock import AsyncMock, patch

import pytest
import fitz
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, RGBColor
from fastapi import FastAPI
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, func, select

import parse_video_py.document_summary as summary_service
import parse_video_py.document_summary_web as summary_web
import parse_video_py.document_translation as translation_service
import parse_video_py.document_translation_web as translation_web
import parse_video_py.user_db as user_db

_SAMPLE_PNG = (Path(__file__).parent / "assets" / "sample_ocr.png").read_bytes()


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend((run_properties, text_node))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_page_field(paragraph) -> None:
    begin = paragraph.add_run()
    begin_node = OxmlElement("w:fldChar")
    begin_node.set(qn("w:fldCharType"), "begin")
    begin._r.append(begin_node)
    instruction = paragraph.add_run()
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = " PAGE "
    instruction._r.append(instruction_node)
    separate = paragraph.add_run()
    separate_node = OxmlElement("w:fldChar")
    separate_node.set(qn("w:fldCharType"), "separate")
    separate._r.append(separate_node)
    paragraph.add_run("1")
    end = paragraph.add_run()
    end_node = OxmlElement("w:fldChar")
    end_node.set(qn("w:fldCharType"), "end")
    end._r.append(end_node)


def _docx_structure_snapshot(document) -> dict:
    story_structure = []
    for section in document.sections:
        for story in (section.header, section.footer):
            story_structure.append(
                {
                    "linked": story.is_linked_to_previous,
                    "paragraph_styles": [
                        paragraph.style.name for paragraph in story.paragraphs
                    ],
                    "table_count": len(story.tables),
                    "drawing_count": len(story._element.xpath(".//w:drawing")),
                    "hyperlink_count": len(story._element.xpath(".//w:hyperlink")),
                    "tab_count": len(story._element.xpath(".//w:tab")),
                    "break_count": len(story._element.xpath(".//w:br")),
                    "field_count": len(story._element.xpath(".//w:fldChar")),
                    "paragraph_properties": [
                        node.xml for node in story._element.xpath(".//w:pPr")
                    ],
                    "run_properties": [
                        node.xml for node in story._element.xpath(".//w:rPr")
                    ],
                    "table_properties": [
                        node.xml for node in story._element.xpath(".//w:tblPr")
                    ],
                }
            )
    return {
        "sections": [
            (
                section.orientation,
                section.page_width,
                section.page_height,
                section.top_margin,
                section.right_margin,
                section.bottom_margin,
                section.left_margin,
            )
            for section in document.sections
        ],
        "paragraph_styles": [paragraph.style.name for paragraph in document.paragraphs],
        "table_count": len(document.tables),
        "inline_shape_count": len(document.inline_shapes),
        "drawing_count": len(document.element.xpath(".//w:drawing")),
        "hyperlink_count": len(document.element.xpath(".//w:hyperlink")),
        "tab_count": len(document.element.xpath(".//w:tab")),
        "break_count": len(document.element.xpath(".//w:br")),
        "bookmark_start_count": len(document.element.xpath(".//w:bookmarkStart")),
        "bookmark_end_count": len(document.element.xpath(".//w:bookmarkEnd")),
        "field_count": len(document.element.xpath(".//w:fldChar")),
        "paragraph_properties": [
            node.xml for node in document.element.xpath(".//w:pPr")
        ],
        "run_properties": [node.xml for node in document.element.xpath(".//w:rPr")],
        "table_properties": [node.xml for node in document.element.xpath(".//w:tblPr")],
        "section_properties": [
            node.xml for node in document.element.xpath(".//w:sectPr")
        ],
        "stories": story_structure,
    }


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("1. Service Agreement", level=1)
    document.add_paragraph("The total amount is USD 1,200. Visit https://example.com.")
    document.add_paragraph("Delivery date: 2026-08-01.")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _many_paragraph_docx_bytes(count: int = 45) -> bytes:
    document = Document()
    for index in range(count):
        document.add_paragraph(f"Paragraph {index} needs translation.")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _styled_docx_bytes() -> bytes:
    document = Document()
    first_section = document.sections[0]
    first_section.page_width = Inches(8.27)
    first_section.page_height = Inches(11.69)
    first_section.top_margin = Inches(0.7)
    first_section.right_margin = Inches(0.8)
    first_section.bottom_margin = Inches(0.9)
    first_section.left_margin = Inches(1.0)
    first_section.header.paragraphs[0].text = "Agreement header"
    first_section.footer.paragraphs[0].text = "Page "
    _add_page_field(first_section.footer.paragraphs[0])
    document.add_heading("1. Service Agreement", level=1)
    paragraph = document.add_paragraph(
        "The total amount is USD 1,200. Visit https://example.com."
    )
    paragraph.style = "Intense Quote"
    paragraph.runs[0].font.bold = True
    paragraph.runs[0].font.italic = True
    paragraph.runs[0].font.color.rgb = RGBColor(0x44, 0x55, 0x66)
    ordinary = document.add_paragraph(style="Normal")
    ordinary.add_run("Short body ").bold = True
    ordinary.add_run("with mixed formatting. ").italic = True
    _add_hyperlink(ordinary, "OpenAI", "https://openai.com")
    ordinary.add_run("\tTabbed")
    ordinary.add_run().add_break()
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), "42")
    bookmark_start.set(qn("w:name"), "translationBookmark")
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), "42")
    ordinary._p.insert(0, bookmark_start)
    ordinary._p.append(bookmark_end)
    ordinary.add_run().add_picture(io.BytesIO(_SAMPLE_PNG), width=Inches(0.2))
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Amount"
    second_section = document.add_section(WD_SECTION.NEW_PAGE)
    second_section.orientation = WD_ORIENT.LANDSCAPE
    second_section.page_width = Inches(11.69)
    second_section.page_height = Inches(8.27)
    second_section.top_margin = Inches(0.6)
    second_section.right_margin = Inches(0.6)
    second_section.bottom_margin = Inches(0.6)
    second_section.left_margin = Inches(0.6)
    second_section.header.is_linked_to_previous = False
    second_section.header.paragraphs[0].text = "Second section header"
    second_section.footer.is_linked_to_previous = False
    second_section.footer.paragraphs[0].text = "Second section footer"
    document.add_paragraph("Second section body", style="List Number")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _styled_pdf_bytes() -> bytes:
    document = fitz.open()
    page = document.new_page(width=420, height=300)
    page.draw_rect(fitz.Rect(24, 24, 396, 276), color=(0.1, 0.3, 0.8), width=2)
    page.insert_textbox(
        fitz.Rect(50, 55, 370, 90),
        "Service Agreement",
        fontsize=18,
        fontname="helv",
        align=fitz.TEXT_ALIGN_CENTER,
        color=(0.1, 0.2, 0.5),
    )
    page.insert_textbox(
        fitz.Rect(55, 120, 365, 175),
        "The total amount is USD 1,200.",
        fontsize=11,
        fontname="helv",
    )
    page.insert_image(fitz.Rect(350, 230, 385, 265), stream=_SAMPLE_PNG)
    payload = document.tobytes(deflate=True)
    document.close()
    return payload


class _FakePdfLayoutPage:
    def __init__(self, fitter, *, images=None, drawings=None):
        self.rect = fitz.Rect(0, 0, 300, 300)
        self._fitter = fitter
        self._images = images or []
        self._drawings = drawings or []
        self.calls = []

    def get_image_info(self, **_kwargs):
        return self._images

    def get_drawings(self):
        return self._drawings

    def insert_htmlbox(self, rect, content, *, css, scale_low, overlay):
        rect = fitz.Rect(rect)
        self.calls.append(
            {
                "rect": rect,
                "content": content,
                "css": css,
                "scale_low": scale_low,
                "overlay": overlay,
            }
        )
        return self._fitter(rect, scale_low)


def _pdf_layout_segment(
    *,
    segment_id="p0001-b0001",
    source_text="原文",
    translated_text="Translated text",
    bbox=(40, 40, 200, 70),
    kind="paragraph",
):
    return {
        "segment_id": segment_id,
        "source_text": source_text,
        "translated_text": translated_text,
        "kind": kind,
        "layout": {
            "page_number": 1,
            "block_index": 1,
            "bbox": list(bbox),
            "font_size": 12,
            "font_color": "#000000",
            "font_family": "sans-serif",
            "text_align": "left",
        },
    }


def _translated_batch(segments, **_kwargs):
    return "en", [
        {
            "segment_id": segment["segment_id"],
            "translated_text": f"译文：{segment['source_text']}",
        }
        for segment in segments
    ]


def _build_app(monkeypatch, tmp_path):
    engine = create_engine(
        f"sqlite+pysqlite:///{(tmp_path / 'translation.db').as_posix()}",
        connect_args={"check_same_thread": False},
        future=True,
    )
    monkeypatch.setattr(user_db, "_engine", engine)
    monkeypatch.setattr(summary_service, "_engine", engine)
    monkeypatch.setattr(translation_service, "_engine", engine)
    monkeypatch.setattr(
        translation_web, "schedule_translation_job", lambda _job_id: None
    )
    monkeypatch.setattr(summary_web, "UPLOAD_DIR", tmp_path / "uploads")
    user_db.init_user_database()
    user = user_db.get_or_create_user("translation-user-openid")
    app = FastAPI()
    app.include_router(summary_web.router)
    app.include_router(translation_web.router)
    return TestClient(app), user, engine


def _complete_translation_job(client, response):
    assert response.status_code == 202
    job_id = response.json()["job_id"]
    asyncio.run(translation_service.run_translation_job(job_id))
    status = client.get(
        f"/auth/documents/translation-jobs/{job_id}",
        cookies={"web_session": "session"},
    )
    assert status.status_code == 200
    assert status.json()["status"] == "completed"
    return status.json()


def _upload_and_parse(client):
    return _upload_bytes_and_parse(client, _docx_bytes(), "agreement.docx")


def _upload_bytes_and_parse(client, content: bytes, filename: str):
    content_type = (
        "application/pdf"
        if filename.lower().endswith(".pdf")
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    uploaded = client.post(
        "/auth/documents/upload",
        files={
            "file": (
                filename,
                content,
                content_type,
            )
        },
        cookies={"web_session": "session"},
    )
    assert uploaded.status_code == 200
    document_id = uploaded.json()["document_id"]
    parsed = client.post(
        f"/auth/documents/{document_id}/parse",
        cookies={"web_session": "session"},
    )
    assert parsed.status_code == 200
    return document_id


def test_translation_cache_modes_glossary_and_exports(monkeypatch, tmp_path):
    client, user, engine = _build_app(monkeypatch, tmp_path)
    with patch.object(summary_web, "verify_web_session", return_value=user):
        document_id = _upload_and_parse(client)
    mocked_ai = AsyncMock(side_effect=_translated_batch)
    payload = {
        "source_language": "auto",
        "target_language": "zh-CN",
        "mode": "bilingual",
        "style": "legal",
        "glossary": [{"source": "Service Agreement", "target": "服务协议"}],
    }
    with (
        patch.object(summary_web, "verify_web_session", return_value=user),
        patch.object(translation_web, "_current_user", return_value=user),
        patch.object(translation_service, "call_deepseek_translation", mocked_ai),
    ):
        first = client.post(
            f"/auth/documents/{document_id}/translate",
            json=payload,
            cookies={"web_session": "session"},
        )
        completed = _complete_translation_job(client, first)
        second = client.post(
            f"/auth/documents/{document_id}/translate",
            json=payload,
            cookies={"web_session": "session"},
        )
        translation_id = completed["result"]["translation_id"]
        docx = client.get(
            f"/auth/documents/{document_id}/translations/{translation_id}/export?format=docx",
            cookies={"web_session": "session"},
        )
        invalid_pdf_export = client.get(
            f"/auth/documents/{document_id}/translations/{translation_id}/export?format=pdf",
            cookies={"web_session": "session"},
        )

    assert first.status_code == 202
    assert completed["result"]["cached"] is False
    assert completed["result"]["detected_source_language"] == "en"
    assert [item["segment_id"] for item in completed["result"]["segments"]] == [
        "body-p0001",
        "body-p0002",
        "body-p0003",
    ]
    assert all(item.get("location") for item in completed["result"]["segments"])
    assert second.status_code == 202
    assert second.json()["job_id"] == first.json()["job_id"]
    assert second.json()["status"] == "completed"
    assert mocked_ai.await_count == 1
    assert docx.status_code == 200
    assert docx.content.startswith(b"PK")
    assert invalid_pdf_export.status_code == 422
    assert invalid_pdf_export.json()["detail"] == "DOCX翻译结果仅支持导出为DOCX"
    with engine.connect() as conn:
        count = conn.execute(
            select(func.count())
            .select_from(user_db.document_translations)
            .where(user_db.document_translations.c.user_id == user.id)
        ).scalar_one()
    assert count == 1


def test_docx_translation_export_preserves_original_layout(monkeypatch, tmp_path):
    client, user, _ = _build_app(monkeypatch, tmp_path)
    source_bytes = _styled_docx_bytes()
    source_document = Document(io.BytesIO(source_bytes))
    source_snapshot = _docx_structure_snapshot(source_document)
    with patch.object(summary_web, "verify_web_session", return_value=user):
        document_id = _upload_bytes_and_parse(
            client, source_bytes, "styled-agreement.docx"
        )
    mocked_ai = AsyncMock(side_effect=_translated_batch)
    with (
        patch.object(summary_web, "verify_web_session", return_value=user),
        patch.object(translation_web, "_current_user", return_value=user),
        patch.object(translation_service, "call_deepseek_translation", mocked_ai),
    ):
        translated = client.post(
            f"/auth/documents/{document_id}/translate",
            json={"target_language": "zh-CN", "mode": "translation"},
            cookies={"web_session": "session"},
        )
        completed = _complete_translation_job(client, translated)
        translation_id = completed["result"]["translation_id"]
        with patch.object(
            translation_service, "Document", wraps=Document
        ) as document_factory:
            exported = client.get(
                f"/auth/documents/{document_id}/translations/{translation_id}/export?format=docx",
                cookies={"web_session": "session"},
            )

    assert exported.status_code == 200
    output = Document(io.BytesIO(exported.content))
    assert _docx_structure_snapshot(output) == source_snapshot
    assert all(call.args and call.args[0] for call in document_factory.call_args_list)
    assert "Document()" not in inspect.getsource(
        translation_service._render_docx_from_original
    )
    assert "add_heading" not in inspect.getsource(
        translation_service._render_docx_from_original
    )
    assert output.paragraphs[0].style.name == "Heading 1"
    assert output.paragraphs[0].text == "译文：1. Service Agreement"
    assert output.paragraphs[1].style.name == "Intense Quote"
    assert output.paragraphs[1].text.startswith("译文：The total amount")
    assert output.paragraphs[2].style.name == "Normal"
    assert output.paragraphs[-1].style.name == "List Number"
    assert output.tables[0].style.name == "Table Grid"
    assert output.tables[0].rows[0].cells[0].text == "译文：Item"
    assert output.tables[0].rows[0].cells[1].text == "译文：Amount"

    result_segments = completed["result"]["segments"]
    assert result_segments
    assert all(segment.get("location") for segment in result_segments)
    assert {segment["location"]["scope"] for segment in result_segments} >= {
        "body",
        "table",
        "header",
        "footer",
    }
    for segment in result_segments:
        paragraph = translation_service._resolve_docx_location(
            output, segment["location"]
        )
        text_nodes = [
            node
            for token_type, node in translation_service._docx_paragraph_tokens(
                paragraph
            )
            if token_type == "text"
        ]
        actual_text = "".join(node.text or "" for node in text_nodes)
        expected_text = segment["translated_text"].replace("\t", "").replace("\n", "")
        assert actual_text == expected_text


def test_pdf_translation_export_preserves_original_page_layout(monkeypatch, tmp_path):
    client, user, _ = _build_app(monkeypatch, tmp_path)
    source_bytes = _styled_pdf_bytes()
    with patch.object(summary_web, "verify_web_session", return_value=user):
        document_id = _upload_bytes_and_parse(client, source_bytes, "agreement.pdf")
    stored_pdf = Path(
        translation_service.get_owned_document(user.id, document_id)["storage_path"]
    )
    stored_source_before_export = stored_pdf.read_bytes()

    async def translated_pdf_batch(segments, **_kwargs):
        translations = ["服务协议", "总金额为 1,200 美元"]
        assert len(segments) == len(translations)
        return "en", [
            {
                "segment_id": segment["segment_id"],
                "translated_text": translated_text,
            }
            for segment, translated_text in zip(segments, translations, strict=False)
        ]

    with (
        patch.object(summary_web, "verify_web_session", return_value=user),
        patch.object(translation_web, "_current_user", return_value=user),
        patch.object(
            translation_service,
            "call_deepseek_translation",
            side_effect=translated_pdf_batch,
        ),
    ):
        translated = client.post(
            f"/auth/documents/{document_id}/translate",
            json={"target_language": "zh-CN", "mode": "translation"},
            cookies={"web_session": "session"},
        )
        completed = _complete_translation_job(client, translated)
        assert translated.status_code == 202
        assert [item["segment_id"] for item in completed["result"]["segments"]] == [
            "p0001-b0001",
            "p0001-b0002",
        ]
        translation_id = completed["result"]["translation_id"]
        exported = client.get(
            f"/auth/documents/{document_id}/translations/{translation_id}/export?format=pdf",
            cookies={"web_session": "session"},
        )
        invalid_docx_export = client.get(
            f"/auth/documents/{document_id}/translations/{translation_id}/export?format=docx",
            cookies={"web_session": "session"},
        )

    assert exported.status_code == 200
    assert invalid_docx_export.status_code == 422
    assert invalid_docx_export.json()["detail"] == "PDF翻译结果仅支持导出为PDF"
    assert exported.headers["content-type"].startswith("application/pdf")
    with fitz.open(stream=source_bytes, filetype="pdf") as source_document:
        source_page = source_document[0]
        source_size = (source_page.rect.width, source_page.rect.height)
        source_drawing_count = len(source_page.get_drawings())
        source_image_count = len(source_page.get_images(full=True))
    with fitz.open(stream=exported.content, filetype="pdf") as output_document:
        assert output_document.page_count == 1
        output_page = output_document[0]
        assert (output_page.rect.width, output_page.rect.height) == source_size
        assert len(output_page.get_drawings()) >= source_drawing_count
        assert len(output_page.get_images(full=True)) == source_image_count
        output_text = output_page.get_text("text")
        assert "Service Agreement" not in output_text
        assert "The total amount is USD 1,200." not in output_text
        assert "服务协议" in output_text
        assert "总金额为 1,200 美元" in output_text
        for segment in completed["result"]["segments"]:
            translated_text = segment["translated_text"]
            bbox_text = output_page.get_textbox(fitz.Rect(segment["layout"]["bbox"]))
            assert translated_text in bbox_text

    with (
        patch.object(translation_web, "_current_user", return_value=user),
        patch.object(
            translation_service,
            "_insert_pdf_translations",
            side_effect=translation_service.DocumentTranslationError(
                "PDF版式写入失败：测试错误"
            ),
        ),
    ):
        failed_export = client.get(
            f"/auth/documents/{document_id}/translations/{translation_id}/export?format=pdf",
            cookies={"web_session": "session"},
        )
    assert failed_export.status_code == 502
    assert failed_export.json()["detail"] == "PDF版式写入失败：测试错误"
    assert stored_pdf.read_bytes() == stored_source_before_export


def test_pdf_physical_line_breaks_are_reflowed_as_natural_paragraphs():
    translated = (
        "This line was wrapped by PDF geometry\n"
        "but belongs to the same paragraph.\n\n"
        "This is an explicit second paragraph."
    )
    normalized = translation_service._normalize_pdf_translated_text(
        translated, "paragraph"
    )
    assert normalized == (
        "This line was wrapped by PDF geometry but belongs to the same paragraph."
        "\n\nThis is an explicit second paragraph."
    )
    assert translation_service._pdf_translation_html(normalized).count("<br>") == 2
    assert (
        translation_service._normalize_pdf_translated_text(
            "A title\nwrapped physically", "heading"
        )
        == "A title wrapped physically"
    )


def test_pdf_long_chinese_to_english_translation_scales_to_point_35(monkeypatch):
    source_text = (
        "这是一段用于验证PDF长文本翻译布局、自动缩放、区域扩展、"
        "页面边界以及障碍物避让的中文内容。"
    )
    translated_text = (
        "This is an English translation used to verify long PDF text layout, "
        "automatic scaling, and safe region expansion."
    )
    assert 2 <= len(translated_text) / len(source_text) <= 3
    segment = _pdf_layout_segment(
        source_text=source_text,
        translated_text=translated_text,
    )
    page = _FakePdfLayoutPage(
        lambda _rect, scale_low: ((0, 0.35) if scale_low <= 0.35 else (-1, scale_low))
    )
    monkeypatch.setattr(translation_service, "PDF_TRANSLATION_MIN_SCALE", 0.35)

    plan = translation_service._fit_pdf_layout_block(page, segment, [segment])

    assert [call["scale_low"] for call in page.calls] == [1.0, 0.35]
    assert plan["scale_low"] == 0.35
    assert plan["fitted_scale"] == 0.35
    assert "\n" not in plan["normalized_text"]


def test_pdf_layout_expands_bbox_down_before_horizontal(monkeypatch):
    segment = _pdf_layout_segment(
        translated_text="A translated paragraph that needs additional vertical space."
    )
    page = _FakePdfLayoutPage(
        lambda rect, _scale: ((0, 0.35) if rect.y1 > 70 else (-1, 1.0))
    )
    monkeypatch.setattr(translation_service, "PDF_TRANSLATION_MIN_SCALE", 0.35)

    plan = translation_service._fit_pdf_layout_block(page, segment, [segment])

    assert plan["bbox"][3] > segment["layout"]["bbox"][3]
    assert page.calls[2]["rect"].y1 > page.calls[1]["rect"].y1


def test_pdf_layout_expansion_stops_before_text_image_and_graphics(monkeypatch):
    segment = _pdf_layout_segment()
    next_segment = _pdf_layout_segment(
        segment_id="p0001-b0002",
        bbox=(40, 120, 200, 145),
    )
    page = _FakePdfLayoutPage(
        lambda _rect, _scale: (-1, 1.0),
        images=[{"bbox": (40, 105, 200, 115)}],
        drawings=[
            {
                "width": 1,
                "items": [("l", fitz.Point(20, 95), fitz.Point(250, 95))],
            }
        ],
    )
    monkeypatch.setattr(translation_service, "PDF_TRANSLATION_MIN_SCALE", 0.35)

    attempts = translation_service._pdf_layout_attempts(
        page, segment, [segment, next_segment]
    )
    expanded_rects = [
        rect for stage, rect, _ in attempts if stage.startswith("expanded")
    ]

    assert expanded_rects
    assert all(rect.y1 <= 92.5 for rect in expanded_rects)
    assert all(
        rect.x0 >= page.rect.x0 and rect.x1 <= page.rect.x1 for rect in expanded_rects
    )


def test_pdf_layout_overflow_has_structured_422(monkeypatch):
    segment = _pdf_layout_segment(
        translated_text="This translation can never fit in the available PDF area."
    )
    page = _FakePdfLayoutPage(lambda _rect, _scale: (-1, 1.0))
    monkeypatch.setattr(translation_service, "PDF_TRANSLATION_MIN_SCALE", 0.35)

    with pytest.raises(translation_service.PDFLayoutOverflowError) as error_info:
        translation_service._fit_pdf_layout_block(page, segment, [segment])

    error = error_info.value
    assert error.error_code == "PDF_LAYOUT_OVERFLOW"
    assert error.page_number == 1
    assert error.segment_id == "p0001-b0001"
    assert error.original_bbox == [40.0, 40.0, 200.0, 70.0]
    assert error.attempted_scale == 0.25
    assert error.translated_length == len(segment["translated_text"])
    assert 0.25 in [call["scale_low"] for call in page.calls]

    response_error = translation_web._service_error(error)
    assert response_error.status_code == 422
    assert response_error.detail["error_code"] == "PDF_LAYOUT_OVERFLOW"
    assert response_error.detail["page_number"] == 1
    assert response_error.detail["attempted_bbox"] == error.attempted_bbox


def test_original_layout_pdf_export_rejects_bilingual_mode(monkeypatch, tmp_path):
    client, user, _ = _build_app(monkeypatch, tmp_path)
    with patch.object(summary_web, "verify_web_session", return_value=user):
        document_id = _upload_bytes_and_parse(
            client, _styled_pdf_bytes(), "agreement.pdf"
        )
    with patch.object(translation_web, "_current_user", return_value=user):
        translated = client.post(
            f"/auth/documents/{document_id}/translate",
            json={"target_language": "zh-CN", "mode": "bilingual"},
            cookies={"web_session": "session"},
        )

    assert translated.status_code == 422
    assert translated.json()["detail"] == "PDF翻译暂不支持双语模式"
def test_old_docx_result_without_location_requires_retranslation(monkeypatch, tmp_path):
    client, user, _ = _build_app(monkeypatch, tmp_path)
    with patch.object(summary_web, "verify_web_session", return_value=user):
        document_id = _upload_and_parse(client)
    document = translation_service.get_owned_document(user.id, document_id)

    with pytest.raises(
        translation_service.InvalidTranslationRequestError,
        match="旧DOCX翻译结果缺少结构位置，请重新翻译后再导出",
    ):
        translation_service._render_docx_from_original(
            document,
            [
                {
                    "segment_id": "seg-0001",
                    "source_text": "Service Agreement",
                    "translated_text": "服务协议",
                }
            ],
        )

    assert translation_service.TRANSLATION_PIPELINE_VERSION == "structure-location-v3"


def test_changed_options_create_separate_cached_result(monkeypatch, tmp_path):
    client, user, _ = _build_app(monkeypatch, tmp_path)
    with patch.object(summary_web, "verify_web_session", return_value=user):
        document_id = _upload_and_parse(client)
    mocked_ai = AsyncMock(side_effect=_translated_batch)
    with (
        patch.object(translation_web, "_current_user", return_value=user),
        patch.object(translation_service, "call_deepseek_translation", mocked_ai),
    ):
        general = client.post(
            f"/auth/documents/{document_id}/translate",
            json={"target_language": "zh-CN", "style": "general"},
            cookies={"web_session": "session"},
        )
        completed_general = _complete_translation_job(client, general)
        technical = client.post(
            f"/auth/documents/{document_id}/translate",
            json={"target_language": "zh-CN", "style": "technical"},
            cookies={"web_session": "session"},
        )
        completed_technical = _complete_translation_job(client, technical)
    assert general.status_code == technical.status_code == 202
    assert (
        completed_general["result"]["translation_id"]
        != completed_technical["result"]["translation_id"]
    )
    assert mocked_ai.await_count == 2


def test_duplicate_job_submission_schedules_only_once(monkeypatch, tmp_path):
    client, user, _ = _build_app(monkeypatch, tmp_path)
    with patch.object(summary_web, "verify_web_session", return_value=user):
        document_id = _upload_and_parse(client)
    scheduled = []
    with (
        patch.object(translation_web, "_current_user", return_value=user),
        patch.object(
            translation_web,
            "schedule_translation_job",
            side_effect=lambda job_id: scheduled.append(job_id),
        ),
    ):
        first = client.post(
            f"/auth/documents/{document_id}/translate",
            json={"target_language": "zh-CN"},
            cookies={"web_session": "session"},
        )
        second = client.post(
            f"/auth/documents/{document_id}/translate",
            json={"target_language": "zh-CN"},
            cookies={"web_session": "session"},
        )

    assert first.status_code == second.status_code == 202
    assert first.json()["job_id"] == second.json()["job_id"]
    assert first.json()["status"] == second.json()["status"] == "pending"
    assert scheduled == [first.json()["job_id"]]


def test_long_translation_runs_after_202_and_polling_survives_refresh(
    monkeypatch, tmp_path
):
    client, user, _ = _build_app(monkeypatch, tmp_path)
    with patch.object(summary_web, "verify_web_session", return_value=user):
        document_id = _upload_and_parse(client)
    started = asyncio.Event()
    release = asyncio.Event()

    async def slow_batch(segments, **_kwargs):
        started.set()
        await release.wait()
        return _translated_batch(segments)

    with (
        patch.object(translation_web, "_current_user", return_value=user),
        patch.object(
            translation_service,
            "call_deepseek_translation",
            side_effect=slow_batch,
        ),
    ):
        created = client.post(
            f"/auth/documents/{document_id}/translate",
            json={"target_language": "zh-CN"},
            cookies={"web_session": "session"},
        )
        assert created.status_code == 202
        job_id = created.json()["job_id"]
        assert created.json()["status"] == "pending"

        async def run_detached_job():
            task = asyncio.create_task(translation_service.run_translation_job(job_id))
            await started.wait()
            processing = translation_service.get_translation_job(user.id, job_id)
            assert processing["status"] == "processing"
            assert processing["progress"]["percent"] == 0
            release.set()
            await task

        asyncio.run(run_detached_job())

        # A new client represents a refreshed browser; job state is database-backed.
        refreshed_client = TestClient(client.app)
        refreshed = refreshed_client.get(
            f"/auth/documents/translation-jobs/{job_id}",
            cookies={"web_session": "session"},
        )

    assert refreshed.status_code == 200
    assert refreshed.json()["status"] == "completed"
    assert refreshed.json()["progress"]["percent"] == 100
    assert refreshed.json()["result"]["translation_id"]


def test_translation_job_reports_batch_progress(monkeypatch, tmp_path):
    client, user, _ = _build_app(monkeypatch, tmp_path)
    with patch.object(summary_web, "verify_web_session", return_value=user):
        document_id = _upload_bytes_and_parse(
            client,
            _many_paragraph_docx_bytes(),
            "many-paragraphs.docx",
        )
    progress_updates = []
    original_progress = translation_service._set_translation_job_progress

    def record_progress(job_id, completed, total):
        progress_updates.append((completed, total))
        original_progress(job_id, completed, total)

    mocked_ai = AsyncMock(side_effect=_translated_batch)
    with (
        patch.object(translation_web, "_current_user", return_value=user),
        patch.object(translation_service, "call_deepseek_translation", mocked_ai),
        patch.object(
            translation_service,
            "_set_translation_job_progress",
            side_effect=record_progress,
        ),
    ):
        created = client.post(
            f"/auth/documents/{document_id}/translate",
            json={"target_language": "zh-CN"},
            cookies={"web_session": "session"},
        )
        completed = _complete_translation_job(client, created)

    assert progress_updates == [(0, 3), (1, 3), (2, 3), (3, 3)]
    assert completed["progress"] == {
        "completed_batches": 3,
        "total_batches": 3,
        "percent": 100,
    }


def test_deepseek_failure_is_persisted_on_job(monkeypatch, tmp_path):
    client, user, _ = _build_app(monkeypatch, tmp_path)
    with patch.object(summary_web, "verify_web_session", return_value=user):
        document_id = _upload_and_parse(client)
    with (
        patch.object(translation_web, "_current_user", return_value=user),
        patch.object(
            translation_service,
            "call_deepseek_translation",
            side_effect=translation_service.DocumentTranslationError(
                "DeepSeek API 网络连接失败，请重试"
            ),
        ),
    ):
        created = client.post(
            f"/auth/documents/{document_id}/translate",
            json={"target_language": "zh-CN"},
            cookies={"web_session": "session"},
        )
        asyncio.run(translation_service.run_translation_job(created.json()["job_id"]))
        failed = client.get(
            f"/auth/documents/translation-jobs/{created.json()['job_id']}",
            cookies={"web_session": "session"},
        )

    assert failed.status_code == 200
    assert failed.json()["status"] == "failed"
    assert failed.json()["error_code"] == "DEEPSEEK_CONNECTION_FAILED"
    assert failed.json()["message"] == "DeepSeek API 网络连接失败，请重试"
    assert failed.json()["request_id"]


def test_translation_job_timeout_and_status_expiry_cleanup(monkeypatch, tmp_path):
    client, user, engine = _build_app(monkeypatch, tmp_path)
    with patch.object(summary_web, "verify_web_session", return_value=user):
        document_id = _upload_and_parse(client)

    async def never_finishes(_segments, **_kwargs):
        await asyncio.sleep(1)

    with (
        patch.object(translation_web, "_current_user", return_value=user),
        patch.object(
            translation_service,
            "call_deepseek_translation",
            side_effect=never_finishes,
        ),
        patch.object(translation_service, "TRANSLATION_JOB_TIMEOUT_SECONDS", 0.01),
    ):
        created = client.post(
            f"/auth/documents/{document_id}/translate",
            json={"target_language": "zh-CN"},
            cookies={"web_session": "session"},
        )
        job_id = created.json()["job_id"]
        asyncio.run(translation_service.run_translation_job(job_id))
        failed = translation_service.get_translation_job(user.id, job_id)

    assert failed["status"] == "failed"
    assert failed["error_code"] == "TRANSLATION_JOB_TIMEOUT"
    with engine.connect() as conn:
        translation_status = conn.execute(
            select(user_db.document_translations.c.status).where(
                user_db.document_translations.c.document_id == document_id
            )
        ).scalar_one()
    assert translation_status == "failed"
    cleanup = translation_service.cleanup_translation_jobs(failed["expires_at"] + 1)
    assert cleanup["deleted"] == 1
    with pytest.raises(KeyError, match="不存在或已过期"):
        translation_service.get_translation_job(user.id, job_id)


def test_translation_requires_parsed_owned_document(monkeypatch, tmp_path):
    client, user, _ = _build_app(monkeypatch, tmp_path)
    with patch.object(summary_web, "verify_web_session", return_value=user):
        uploaded = client.post(
            "/auth/documents/upload",
            files={"file": ("agreement.docx", _docx_bytes())},
            cookies={"web_session": "session"},
        )
    with patch.object(translation_web, "_current_user", return_value=user):
        unparsed = client.post(
            f"/auth/documents/{uploaded.json()['document_id']}/translate",
            json={"target_language": "zh-CN"},
            cookies={"web_session": "session"},
        )
    assert unparsed.status_code == 502
    assert unparsed.json()["detail"] == "请先完成文档解析"

    other = user_db.get_or_create_user("other-translation-user")
    with patch.object(translation_web, "_current_user", return_value=other):
        forbidden = client.post(
            f"/auth/documents/{uploaded.json()['document_id']}/translate",
            json={"target_language": "zh-CN"},
            cookies={"web_session": "other"},
        )
    assert forbidden.status_code == 404


def test_segmentation_and_batching_uses_configured_segment_count():
    segments = translation_service.segment_document(
        "Heading\nFirst sentence. Second sentence.\nThird paragraph."
    )
    assert len(segments) == 3
    assert segments[1]["source_text"] == "First sentence. Second sentence."
    batches = list(translation_service.iter_segment_batches(segments, batch_size=2))
    assert [len(batch) for batch in batches] == [2, 1]
    assert [item for batch in batches for item in batch] == segments

    many_segments = translation_service.segment_document(
        "\n".join(f"Paragraph {index}." for index in range(45))
    )
    default_batches = list(translation_service.iter_segment_batches(many_segments))
    assert [len(batch) for batch in default_batches] == [20, 20, 5]


async def test_blank_numeric_url_and_symbol_segments_bypass_model():
    segments = [
        {"segment_id": "blank", "source_text": "   ", "kind": "paragraph"},
        {"segment_id": "number", "source_text": "1,200.00", "kind": "paragraph"},
        {
            "segment_id": "url",
            "source_text": "https://example.com/path?a=1",
            "kind": "paragraph",
        },
        {"segment_id": "symbols", "source_text": "※ → ★ !!!", "kind": "paragraph"},
        {
            "segment_id": "text",
            "source_text": "Service Agreement",
            "kind": "heading",
        },
    ]
    requested_ids = []

    async def translate_text_only(requested, **_kwargs):
        requested_ids.append([item["segment_id"] for item in requested])
        return "en", [
            {"segment_id": "text", "translated_text": "服务协议"},
        ]

    with patch.object(
        translation_service,
        "call_deepseek_translation",
        side_effect=translate_text_only,
    ):
        _, translated = await translation_service.translate_batch_with_recovery(
            segments,
            source_language="auto",
            target_language="zh-CN",
            style="general",
            glossary=[],
            user_id="user-id",
        )

    assert requested_ids == [["text"]]
    assert translated == [
        {"segment_id": "number", "translated_text": "1,200.00"},
        {
            "segment_id": "url",
            "translated_text": "https://example.com/path?a=1",
        },
        {"segment_id": "symbols", "translated_text": "※ → ★ !!!"},
        {"segment_id": "text", "translated_text": "服务协议"},
    ]


def test_segment_document_filters_empty_paragraphs():
    segments = translation_service.segment_document(
        "\n  \nFirst paragraph.\n\t\nSecond."
    )
    assert [item["source_text"] for item in segments] == [
        "First paragraph.",
        "Second.",
    ]


def test_missing_and_duplicate_ids_have_structured_diagnostics():
    segments = translation_service.segment_document("First.\nSecond.")
    with pytest.raises(
        translation_service.IncompleteTranslationResponseError
    ) as missing_error:
        translation_service._normalize_ai_result(
            {
                "translations": [
                    {
                        "segment_id": "seg-0001",
                        "translated_text": "第一段。",
                    }
                ]
            },
            segments,
        )
    assert missing_error.value.expected_count == 2
    assert missing_error.value.actual_count == 1
    assert missing_error.value.missing_ids == ["seg-0002"]
    assert missing_error.value.extra_ids == []
    assert missing_error.value.duplicate_ids == []

    with pytest.raises(
        translation_service.IncompleteTranslationResponseError
    ) as duplicate_error:
        translation_service._normalize_ai_result(
            {
                "translations": [
                    {
                        "segment_id": "seg-0001",
                        "translated_text": "第一段。",
                    },
                    {
                        "segment_id": "seg-0001",
                        "translated_text": "重复第一段。",
                    },
                    {
                        "segment_id": "seg-extra",
                        "translated_text": "额外段落。",
                    },
                ]
            },
            segments,
        )
    assert duplicate_error.value.expected_count == 2
    assert duplicate_error.value.actual_count == 3
    assert duplicate_error.value.missing_ids == ["seg-0002"]
    assert duplicate_error.value.extra_ids == ["seg-extra"]
    assert duplicate_error.value.duplicate_ids == ["seg-0001"]


async def test_deepseek_translation_uses_segment_json_and_glossary(monkeypatch):
    captured = {}
    segments = translation_service.segment_document(
        "Service Agreement\nThe total amount is USD 1,200."
    )

    class FakeResponse:
        status_code = 200

        def raise_for_status(self):
            return None

        def json(self):
            return {
                "choices": [
                    {
                        "message": {
                            "content": json.dumps(
                                {
                                    "translations": [
                                        {
                                            "segment_id": item["segment_id"],
                                            "translated_text": f"译文 {index}",
                                        }
                                        for index, item in enumerate(segments, 1)
                                    ],
                                }
                            )
                        }
                    }
                ]
            }

    class FakeClient:
        async def __aenter__(self):
            return self

        async def __aexit__(self, *_args):
            return None

        async def post(self, url, headers, json):
            captured.update({"url": url, "headers": headers, "body": json})
            return FakeResponse()

    monkeypatch.setenv("DEEPSEEK_API_KEY", "secret")
    with patch.object(
        translation_service.httpx, "AsyncClient", return_value=FakeClient()
    ):
        detected, translated = await translation_service.call_deepseek_translation(
            segments,
            source_language="auto",
            target_language="zh-CN",
            style="business",
            glossary=[{"source": "Service Agreement", "target": "服务协议"}],
            user_id="user-id",
        )

    assert detected is None
    assert [item["segment_id"] for item in translated] == ["seg-0001", "seg-0002"]
    assert captured["headers"]["Authorization"] == "Bearer secret"
    assert captured["body"]["response_format"] == {"type": "json_object"}
    prompt = captured["body"]["messages"][1]["content"]
    assert "Service Agreement" in prompt
    assert "服务协议" in prompt
    assert "USD 1,200" in prompt
    assert "detected_source_language" not in prompt
    assert "禁止遗漏、合并、改名、重复、新增或重排" in prompt
    assert "Markdown" in prompt


async def test_incomplete_response_logs_counts_and_ids(monkeypatch, caplog):
    segments = translation_service.segment_document("First.\nSecond.")

    class FakeResponse:
        status_code = 200

        def raise_for_status(self):
            return None

        def json(self):
            return {
                "choices": [
                    {
                        "finish_reason": "stop",
                        "message": {
                            "content": json.dumps(
                                {
                                    "translations": [
                                        {
                                            "segment_id": "seg-0001",
                                            "translated_text": "第一段。",
                                        },
                                        {
                                            "segment_id": "seg-extra",
                                            "translated_text": "额外段落。",
                                        },
                                        {
                                            "segment_id": "seg-0001",
                                            "translated_text": "重复段落。",
                                        },
                                    ]
                                }
                            )
                        },
                    }
                ]
            }

    class FakeClient:
        async def __aenter__(self):
            return self

        async def __aexit__(self, *_args):
            return None

        async def post(self, *_args, **_kwargs):
            return FakeResponse()

    monkeypatch.setenv("DEEPSEEK_API_KEY", "secret")
    caplog.set_level(logging.WARNING, logger=translation_service.__name__)
    with (
        patch.object(
            translation_service.httpx, "AsyncClient", return_value=FakeClient()
        ),
        pytest.raises(translation_service.IncompleteTranslationResponseError),
    ):
        await translation_service.call_deepseek_translation(
            segments,
            source_language="auto",
            target_language="zh-CN",
            style="general",
            glossary=[],
            user_id="user-id",
            request_id="request-id",
            batch_index=1,
        )

    log_text = caplog.text
    assert "expected_count=2" in log_text
    assert "actual_count=3" in log_text
    assert 'missing_ids=["seg-0002"]' in log_text
    assert 'extra_ids=["seg-extra"]' in log_text
    assert 'duplicate_ids=["seg-0001"]' in log_text
    assert "First." not in log_text
    assert "第一段。" not in log_text


async def test_invalid_json_is_logged_without_response_content(monkeypatch, caplog):
    segments = translation_service.segment_document("Service Agreement")
    invalid_content = "```json not valid ```"

    class FakeResponse:
        status_code = 200

        def raise_for_status(self):
            return None

        def json(self):
            return {
                "choices": [
                    {
                        "finish_reason": "stop",
                        "message": {"content": invalid_content},
                    }
                ]
            }

    class FakeClient:
        async def __aenter__(self):
            return self

        async def __aexit__(self, *_args):
            return None

        async def post(self, *_args, **_kwargs):
            return FakeResponse()

    monkeypatch.setenv("DEEPSEEK_API_KEY", "secret")
    caplog.set_level(logging.WARNING, logger=translation_service.__name__)
    with (
        patch.object(
            translation_service.httpx, "AsyncClient", return_value=FakeClient()
        ),
        pytest.raises(translation_service.InvalidTranslationResponseError),
    ):
        await translation_service.call_deepseek_translation(
            segments,
            source_language="auto",
            target_language="zh-CN",
            style="general",
            glossary=[],
            user_id="user-id",
            request_id="request-id",
        )

    assert "reason=invalid_json" in caplog.text
    assert "json_error_position=" in caplog.text
    assert invalid_content not in caplog.text


async def test_incomplete_translation_retries_current_batch():
    segments = translation_service.segment_document(
        "Heading\nFirst paragraph.\nSecond paragraph."
    )
    requested_ids = []

    async def incomplete_once(requested, **_kwargs):
        requested_ids.append([item["segment_id"] for item in requested])
        if len(requested_ids) == 1:
            payload = {
                "translations": [
                    {
                        "segment_id": requested[0]["segment_id"],
                        "translated_text": "标题",
                    },
                    {
                        "segment_id": requested[1]["segment_id"],
                        "translated_text": "第一段",
                    },
                ],
            }
            return translation_service._normalize_ai_result(payload, requested)
        return None, [
            {
                "segment_id": item["segment_id"],
                "translated_text": f"译文 {item['segment_id']}",
            }
            for item in requested
        ]

    with patch.object(
        translation_service,
        "call_deepseek_translation",
        side_effect=incomplete_once,
    ):
        detected, translated = await translation_service.translate_batch_with_recovery(
            segments,
            source_language="auto",
            target_language="zh-CN",
            style="general",
            glossary=[],
            user_id="user-id",
        )

    assert detected is None
    assert requested_ids == [
        ["seg-0001", "seg-0002", "seg-0003"],
        ["seg-0001", "seg-0002", "seg-0003"],
    ]
    assert [item["translated_text"] for item in translated] == [
        "译文 seg-0001",
        "译文 seg-0002",
        "译文 seg-0003",
    ]


async def test_duplicate_id_response_splits_batch_and_retries():
    segments = translation_service.segment_document("Heading\nFirst paragraph.")
    requested_sizes = []

    async def invalid_batch(requested, **_kwargs):
        requested_sizes.append(len(requested))
        if len(requested) > 1:
            return translation_service._normalize_ai_result(
                {
                    "translations": [
                        {
                            "segment_id": requested[0]["segment_id"],
                            "translated_text": "重复一",
                        },
                        {
                            "segment_id": requested[0]["segment_id"],
                            "translated_text": "重复二",
                        },
                    ]
                },
                requested,
            )
        return "en", [
            {
                "segment_id": requested[0]["segment_id"],
                "translated_text": f"译文 {requested[0]['segment_id']}",
            }
        ]

    with patch.object(
        translation_service,
        "call_deepseek_translation",
        side_effect=invalid_batch,
    ):
        _, translated = await translation_service.translate_batch_with_recovery(
            segments,
            source_language="auto",
            target_language="zh-CN",
            style="general",
            glossary=[],
            user_id="user-id",
        )

    assert requested_sizes == [2, 2, 2, 1, 1]
    assert [item["segment_id"] for item in translated] == ["seg-0001", "seg-0002"]


async def test_single_incomplete_segment_has_bounded_retries(caplog):
    segments = translation_service.segment_document("Only one paragraph.")
    mocked_ai = AsyncMock(
        side_effect=translation_service.IncompleteTranslationResponseError([], "en")
    )
    caplog.set_level(logging.ERROR, logger=translation_service.__name__)
    with (
        patch.object(translation_service, "call_deepseek_translation", mocked_ai),
        pytest.raises(
            translation_service.DocumentTranslationError,
            match="seg-0001",
        ),
    ):
        await translation_service.translate_batch_with_recovery(
            segments,
            source_language="auto",
            target_language="zh-CN",
            style="general",
            glossary=[],
            user_id="user-id",
        )

    assert mocked_ai.await_count == translation_service.TRANSLATION_BATCH_RETRIES + 1
    assert "failed_segment_id=seg-0001" in caplog.text


async def test_incomplete_twenty_segment_batch_bisects_to_singletons_in_order():
    segments = translation_service.segment_document(
        "\n".join(f"Paragraph {index}." for index in range(20))
    )
    requested_batches = []

    async def incomplete_until_single(requested, **_kwargs):
        requested_batches.append(tuple(item["segment_id"] for item in requested))
        if len(requested) > 1:
            return translation_service._normalize_ai_result(
                {
                    "translations": [
                        {
                            "segment_id": item["segment_id"],
                            "translated_text": f"译文 {item['segment_id']}",
                        }
                        for item in requested[:-1]
                    ]
                },
                requested,
            )
        return None, [
            {
                "segment_id": requested[0]["segment_id"],
                "translated_text": f"译文 {requested[0]['segment_id']}",
            }
        ]

    with patch.object(
        translation_service,
        "call_deepseek_translation",
        side_effect=incomplete_until_single,
    ):
        _, translated = await translation_service.translate_batch_with_recovery(
            segments,
            source_language="auto",
            target_language="zh-CN",
            style="general",
            glossary=[],
            user_id="user-id",
        )

    requested_sizes = [len(batch) for batch in requested_batches]
    for expected_size in (20, 10, 5, 2, 1):
        assert expected_size in requested_sizes
    assert max(Counter(requested_batches).values()) == 3
    assert [item["segment_id"] for item in translated] == [
        item["segment_id"] for item in segments
    ]


async def test_length_finish_reason_is_retried_then_split_in_original_order():
    segments = translation_service.segment_document("Heading\nFirst.\nSecond.")
    requested_ids = []

    async def truncated_batch(requested, **_kwargs):
        requested_ids.append([item["segment_id"] for item in requested])
        if len(requested) > 1:
            raise translation_service.TruncatedTranslationResponseError(
                "AI 翻译输出被截断", finish_reason="length"
            )
        return None, [
            {
                "segment_id": requested[0]["segment_id"],
                "translated_text": f"译文 {requested[0]['segment_id']}",
            }
        ]

    with patch.object(
        translation_service,
        "call_deepseek_translation",
        side_effect=truncated_batch,
    ):
        _, translated = await translation_service.translate_batch_with_recovery(
            segments,
            source_language="auto",
            target_language="zh-CN",
            style="general",
            glossary=[],
            user_id="user-id",
            request_id="request-id",
            batch_index=1,
        )

    assert requested_ids[:3] == [["seg-0001", "seg-0002", "seg-0003"]] * 3
    assert [item["segment_id"] for item in translated] == [
        "seg-0001",
        "seg-0002",
        "seg-0003",
    ]
