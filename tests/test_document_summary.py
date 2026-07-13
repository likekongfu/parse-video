import io
from types import SimpleNamespace
from unittest.mock import AsyncMock, patch

import fitz
import httpx
import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, func, select

import parse_video_py.document_summary as summary_service
import parse_video_py.document_summary_web as summary_web
import parse_video_py.user_db as user_db


@pytest.fixture()
def summary_app(monkeypatch, tmp_path):
    engine = create_engine(
        f"sqlite+pysqlite:///{(tmp_path / 'summary.db').as_posix()}",
        connect_args={"check_same_thread": False}, future=True,
    )
    monkeypatch.setattr(user_db, "_engine", engine)
    monkeypatch.setattr(summary_service, "_engine", engine)
    monkeypatch.setattr(summary_web, "UPLOAD_DIR", tmp_path / "uploads")
    user_db.init_user_database()
    user = user_db.get_or_create_user("summary-user-openid")
    app = FastAPI()
    app.include_router(summary_web.router)
    with patch.object(summary_web, "verify_web_session", return_value=user):
        yield TestClient(app), user, engine


def _pdf_bytes(text: str = "合同总金额 100 万元，负责人张三，截止日期 2026-08-01。") -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text, fontname="china-s")
    payload = document.tobytes()
    document.close()
    return payload


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("项目负责人李四")
    document.add_paragraph("交付日期为 2026-09-01")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def test_pdf_upload_parse_summary_cache_and_history(summary_app):
    client, user, engine = summary_app
    pdf = _pdf_bytes()
    uploaded = client.post(
        "/auth/documents/upload",
        files={"file": ("contract.pdf", pdf, "application/pdf")},
        cookies={"web_session": "session"},
    )
    assert uploaded.status_code == 200
    document_id = uploaded.json()["document_id"]

    duplicate = client.post(
        "/auth/documents/upload",
        files={"file": ("duplicate.pdf", pdf, "application/pdf")},
        cookies={"web_session": "session"},
    )
    assert duplicate.json()["document_id"] == document_id
    assert duplicate.json()["cached"] is True

    parsed = client.post(
        f"/auth/documents/{document_id}/parse",
        cookies={"web_session": "session"},
    )
    assert parsed.status_code == 200
    assert parsed.json()["text_length"] > 0
    assert client.post(
        f"/auth/documents/{document_id}/parse",
        cookies={"web_session": "session"},
    ).json()["cached"] is True

    ai_result = {
        "summary": "这是一份合同摘要。",
        "key_points": ["需要按期交付"],
        "people": ["张三"],
        "dates": ["2026-08-01"],
        "amounts": ["100 万元"],
        "risks": ["延期风险"],
    }
    mocked_ai = AsyncMock(return_value=ai_result)
    with patch.object(summary_service, "call_deepseek", mocked_ai):
        first = client.post(
            f"/auth/documents/{document_id}/summarize",
            cookies={"web_session": "session"},
        )
        second = client.post(
            f"/auth/documents/{document_id}/summarize",
            cookies={"web_session": "session"},
        )
    assert first.status_code == 200
    assert first.json()["summary"] == ai_result["summary"]
    assert second.json()["cached"] is True
    assert mocked_ai.await_count == 1

    saved = client.post(
        f"/auth/documents/{document_id}/save",
        cookies={"web_session": "session"},
    )
    assert saved.status_code == 200
    history = client.get(
        "/auth/documents/history", cookies={"web_session": "session"}
    ).json()["documents"]
    assert history[0]["document_id"] == document_id
    assert history[0]["summary"]["people"] == ["张三"]

    with engine.connect() as conn:
        task_count = conn.execute(select(func.count()).select_from(
            user_db.document_tasks
        ).where(user_db.document_tasks.c.user_id == user.id)).scalar_one()
    assert task_count == 2


def test_docx_upload_and_parse(summary_app):
    client, _, _ = summary_app
    uploaded = client.post(
        "/auth/documents/upload",
        files={"file": ("plan.docx", _docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        cookies={"web_session": "session"},
    )
    assert uploaded.status_code == 200
    parsed = client.post(
        f"/auth/documents/{uploaded.json()['document_id']}/parse",
        cookies={"web_session": "session"},
    )
    assert parsed.status_code == 200
    assert parsed.json()["text_length"] > 0


def test_empty_pdf_suggests_ocr_and_failed_parse_can_retry(summary_app):
    client, _, _ = summary_app
    uploaded = client.post(
        "/auth/documents/upload",
        files={"file": ("scan.pdf", _pdf_bytes(""), "application/pdf")},
        cookies={"web_session": "session"},
    )
    document_id = uploaded.json()["document_id"]
    first = client.post(
        f"/auth/documents/{document_id}/parse",
        cookies={"web_session": "session"},
    )
    second = client.post(
        f"/auth/documents/{document_id}/parse",
        cookies={"web_session": "session"},
    )
    assert first.status_code == 422
    assert second.status_code == 422
    assert first.json()["detail"] == "可能为扫描件，请使用 OCR"


def test_upload_rejects_unauthenticated_invalid_and_oversized_files(summary_app):
    client, _, _ = summary_app
    with patch.object(summary_web, "verify_web_session", side_effect=ValueError("bad")):
        unauthorized = client.post(
            "/auth/documents/upload",
            files={"file": ("test.pdf", _pdf_bytes(), "application/pdf")},
            cookies={"web_session": "bad"},
        )
    assert unauthorized.status_code == 401

    invalid = client.post(
        "/auth/documents/upload",
        files={"file": ("test.txt", b"text", "text/plain")},
        cookies={"web_session": "session"},
    )
    assert invalid.status_code == 400

    oversized = client.post(
        "/auth/documents/upload",
        files={"file": ("large.pdf", b"%PDF-" + b"x" * (5 * 1024 * 1024), "application/pdf")},
        cookies={"web_session": "session"},
    )
    assert oversized.status_code == 413


def test_other_user_cannot_access_document(summary_app):
    client, _, _ = summary_app
    uploaded = client.post(
        "/auth/documents/upload",
        files={"file": ("private.pdf", _pdf_bytes(), "application/pdf")},
        cookies={"web_session": "session"},
    )
    other = user_db.get_or_create_user("another-summary-user")
    with patch.object(summary_web, "verify_web_session", return_value=other):
        response = client.post(
            f"/auth/documents/{uploaded.json()['document_id']}/parse",
            cookies={"web_session": "other-session"},
        )
    assert response.status_code == 404


@pytest.mark.asyncio
async def test_deepseek_uses_json_output_and_validates_structured_result(monkeypatch):
    captured = {}

    class FakeResponse:
        status_code = 200

        def raise_for_status(self):
            return None

        def json(self):
            return {
                "choices": [{"message": {"content": (
                    '{"summary":"摘要","key_points":["要点"],"people":[],"dates":[],'
                    '"amounts":[],"risks":[]}'
                )}}]
            }

    class FakeClient:
        async def __aenter__(self):
            return self

        async def __aexit__(self, *args):
            return None

        async def post(self, url, headers, json):
            captured.update({"url": url, "headers": headers, "json": json})
            return FakeResponse()

    monkeypatch.setenv("DEEPSEEK_API_KEY", "secret-key")
    with patch.object(summary_service.httpx, "AsyncClient", return_value=FakeClient()):
        result = await summary_service.call_deepseek("文档正文", "system-user-id")

    assert result["summary"] == "摘要"
    assert captured["headers"]["Authorization"] == "Bearer secret-key"
    assert captured["json"]["response_format"] == {"type": "json_object"}
    assert captured["json"]["user_id"] == "system-user-id"
    assert captured["json"]["model"] == "deepseek-v4-flash"


@pytest.mark.asyncio
async def test_deepseek_network_failure_is_retryable(monkeypatch):
    class FailingClient:
        async def __aenter__(self):
            return self

        async def __aexit__(self, *args):
            return None

        async def post(self, *args, **kwargs):
            raise httpx.ConnectError("offline")

    monkeypatch.setenv("DEEPSEEK_API_KEY", "secret-key")
    with patch.object(summary_service.httpx, "AsyncClient", return_value=FailingClient()):
        with pytest.raises(summary_service.DocumentSummaryError, match="网络连接失败"):
            await summary_service.call_deepseek("文档正文", "system-user-id")
