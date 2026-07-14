"""Persistent PDF/DOCX extraction and DeepSeek summary service."""

from __future__ import annotations

import hashlib
import json
import os
import time
import uuid
from pathlib import Path
from typing import Any

import httpx
from sqlalchemy import and_, insert, or_, select, update
from sqlalchemy.exc import IntegrityError

from parse_video_py.user_db import (
    _engine,
    document_tasks,
    documents,
    init_user_database,
)

MAX_UPLOAD_BYTES = 5 * 1024 * 1024
PROCESSING_STALE_SECONDS = 10 * 60
DEEPSEEK_API_URL = os.getenv(
    "DEEPSEEK_API_URL", "https://api.deepseek.com/chat/completions"
).strip()
DEEPSEEK_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-v4-flash").strip()
DEEPSEEK_TIMEOUT_SECONDS = float(os.getenv("DEEPSEEK_TIMEOUT_SECONDS", "120"))
DEEPSEEK_MAX_INPUT_CHARS = int(os.getenv("DEEPSEEK_MAX_INPUT_CHARS", "100000"))

DOCUMENT_TYPES = (
    "resume",
    "contract",
    "legal",
    "research",
    "financial",
    "meeting",
    "work_report",
    "project",
    "prd",
    "technical",
    "manual",
    "tender",
    "policy",
    "news",
    "study",
    "invoice",
    "lease",
    "medical",
    "purchase_order",
    "general",
)
DOCUMENT_TYPE_LABELS = {
    "resume": "简历",
    "contract": "合同",
    "legal": "法律文书",
    "research": "研究报告",
    "financial": "财务报告",
    "meeting": "会议纪要",
    "work_report": "工作报告",
    "project": "项目文档",
    "prd": "产品需求文档",
    "technical": "技术文档",
    "manual": "操作手册",
    "tender": "招投标文件",
    "policy": "政策文件",
    "news": "新闻稿",
    "study": "学习资料",
    "invoice": "发票",
    "lease": "租赁协议",
    "medical": "医疗文档",
    "purchase_order": "采购订单",
    "general": "通用文档",
}
SUMMARY_SECTION_TEMPLATES = {
    "resume": ["核心优势", "工作经历", "教育背景", "专业技能", "岗位匹配", "关注事项"],
    "contract": [
        "合同主体",
        "标的与范围",
        "权利与义务",
        "金额与付款",
        "期限与节点",
        "风险条款",
    ],
    "legal": [
        "案件与事实",
        "争议焦点",
        "法律依据",
        "证据材料",
        "程序与期限",
        "风险提示",
    ],
    "research": [
        "研究问题",
        "方法与样本",
        "核心发现",
        "关键数据",
        "局限性",
        "研究结论",
    ],
    "financial": [
        "财务概览",
        "收入与成本",
        "盈利与现金流",
        "关键指标",
        "异常事项",
        "风险提示",
    ],
    "meeting": ["会议主题", "参会人员", "核心讨论", "决策结论", "待办事项", "时间节点"],
    "work_report": [
        "工作概览",
        "完成事项",
        "关键成果",
        "问题与风险",
        "后续计划",
        "所需支持",
    ],
    "project": [
        "项目目标",
        "范围与交付物",
        "里程碑",
        "角色分工",
        "预算与资源",
        "风险与依赖",
    ],
    "prd": ["产品目标", "用户与场景", "功能需求", "业务规则", "验收标准", "风险与依赖"],
    "technical": [
        "技术目标",
        "系统架构",
        "关键模块",
        "接口与数据",
        "部署运维",
        "技术风险",
    ],
    "manual": ["适用范围", "准备条件", "操作步骤", "参数说明", "故障处理", "安全注意"],
    "tender": ["招标范围", "资质要求", "技术要求", "商务要求", "评分标准", "关键日期"],
    "policy": [
        "政策目标",
        "适用对象",
        "核心措施",
        "执行要求",
        "时间范围",
        "影响与风险",
    ],
    "news": ["事件概述", "关键人物", "时间地点", "核心事实", "影响与进展", "信息来源"],
    "study": ["学习主题", "核心概念", "重点知识", "案例与公式", "易错点", "复习计划"],
    "invoice": [
        "票据信息",
        "购销双方",
        "商品与服务",
        "金额与税额",
        "日期与编号",
        "异常风险",
    ],
    "lease": [
        "租赁主体",
        "租赁标的",
        "租期与交付",
        "租金与押金",
        "双方义务",
        "违约风险",
    ],
    "medical": [
        "基本信息",
        "主要症状",
        "检查结果",
        "诊断意见",
        "治疗与用药",
        "医疗风险",
    ],
    "purchase_order": [
        "订单信息",
        "供需双方",
        "商品明细",
        "数量与金额",
        "交付与付款",
        "异常风险",
    ],
    "general": [
        "文档概览",
        "核心要点",
        "关键数据",
        "相关人员",
        "时间节点",
        "风险与待办",
    ],
}
RISK_NOTICES = {
    "medical": "医疗内容仅供信息整理，不能替代执业医师的诊断、处方或治疗建议。",
    "legal": "法律内容仅供信息整理，不构成正式法律意见，请由专业法律人士复核。",
    "financial": "财务内容仅供信息整理，不构成投资、审计或税务建议，请由专业人士复核。",
}


class DocumentSummaryError(RuntimeError):
    pass


class DocumentBusyError(DocumentSummaryError):
    pass


class EmptyDocumentTextError(DocumentSummaryError):
    pass


def _document_row(user_id: str, document_id: str):
    init_user_database()
    with _engine.connect() as conn:
        return (
            conn.execute(
                select(documents).where(
                    documents.c.id == document_id,
                    documents.c.user_id == user_id,
                )
            )
            .mappings()
            .first()
        )


def get_owned_document(user_id: str, document_id: str):
    row = _document_row(user_id, document_id)
    if not row:
        raise KeyError("文档不存在")
    return row


def register_document(
    *,
    user_id: str,
    filename: str,
    file_type: str,
    file_size: int,
    content_hash: bytes,
    temporary_path: Path,
    upload_dir: Path,
) -> tuple[dict[str, Any], bool]:
    """Persist an upload, reusing an existing document for identical user content."""
    init_user_database()
    with _engine.connect() as conn:
        existing = (
            conn.execute(
                select(documents).where(
                    documents.c.user_id == user_id,
                    documents.c.content_hash == content_hash,
                )
            )
            .mappings()
            .first()
        )
    if existing:
        temporary_path.unlink(missing_ok=True)
        return dict(existing), True

    now = int(time.time())
    document_id = str(uuid.uuid4())
    upload_dir.mkdir(parents=True, exist_ok=True)
    final_path = upload_dir / f"{document_id}.{file_type}"
    temporary_path.replace(final_path)
    values = {
        "id": document_id,
        "user_id": user_id,
        "filename": Path(filename).name[:255] or f"document.{file_type}",
        "file_type": file_type,
        "file_size": file_size,
        "content_hash": content_hash,
        "storage_path": str(final_path),
        "extracted_text": None,
        "extraction_status": "pending",
        "document_type": None,
        "document_type_source": None,
        "summary_json": None,
        "summary_status": "pending",
        "error_message": None,
        "saved_at": None,
        "created_at": now,
        "updated_at": now,
    }
    try:
        with _engine.begin() as conn:
            conn.execute(insert(documents).values(**values))
    except IntegrityError:
        final_path.unlink(missing_ok=True)
        with _engine.connect() as conn:
            existing = (
                conn.execute(
                    select(documents).where(
                        documents.c.user_id == user_id,
                        documents.c.content_hash == content_hash,
                    )
                )
                .mappings()
                .one()
            )
        return dict(existing), True
    except Exception:
        final_path.unlink(missing_ok=True)
        raise
    return values, False


def sha256_file(path: Path) -> bytes:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.digest()


def _set_task(
    conn,
    document_id: str,
    user_id: str,
    task_type: str,
    status: str,
    error_message: str | None = None,
) -> None:
    now = int(time.time())
    existing_id = conn.execute(
        select(document_tasks.c.id).where(
            document_tasks.c.document_id == document_id,
            document_tasks.c.task_type == task_type,
        )
    ).scalar_one_or_none()
    values = {
        "status": status,
        "error_message": error_message[:1000] if error_message else None,
        "updated_at": now,
        "completed_at": now if status == "completed" else None,
    }
    if existing_id:
        conn.execute(
            update(document_tasks)
            .where(document_tasks.c.id == existing_id)
            .values(**values)
        )
    else:
        conn.execute(
            insert(document_tasks).values(
                id=str(uuid.uuid4()),
                document_id=document_id,
                user_id=user_id,
                task_type=task_type,
                created_at=now,
                **values,
            )
        )


def _acquire(document_id: str, user_id: str, kind: str, *, force: bool = False):
    status_column = (
        documents.c.extraction_status if kind == "parse" else documents.c.summary_status
    )
    result_column = (
        documents.c.extracted_text if kind == "parse" else documents.c.summary_json
    )
    now = int(time.time())
    with _engine.begin() as conn:
        row = (
            conn.execute(
                select(documents).where(
                    documents.c.id == document_id,
                    documents.c.user_id == user_id,
                )
            )
            .mappings()
            .first()
        )
        if not row:
            raise KeyError("文档不存在")
        if (
            not force
            and row[status_column.name] == "completed"
            and row[result_column.name]
        ):
            return dict(row), True
        available_statuses = ["pending", "failed"]
        if force:
            available_statuses.append("completed")
        acquired = conn.execute(
            update(documents)
            .where(
                documents.c.id == document_id,
                documents.c.user_id == user_id,
                or_(
                    status_column.in_(available_statuses),
                    and_(
                        status_column == "processing",
                        documents.c.updated_at < now - PROCESSING_STALE_SECONDS,
                    ),
                ),
            )
            .values(
                **{
                    status_column.name: "processing",
                    "error_message": None,
                    "updated_at": now,
                }
            )
        )
        if acquired.rowcount != 1:
            raise DocumentBusyError("文档正在处理中，请稍后重试")
        _set_task(conn, document_id, user_id, kind, "processing")
    return dict(row), False


def _finish(
    document_id: str,
    user_id: str,
    kind: str,
    *,
    result: str | None = None,
    error: str | None = None,
    document_type: str | None = None,
    document_type_source: str | None = None,
) -> None:
    status_column = (
        documents.c.extraction_status if kind == "parse" else documents.c.summary_status
    )
    result_column = (
        documents.c.extracted_text if kind == "parse" else documents.c.summary_json
    )
    status_value = "failed" if error else "completed"
    values: dict[str, Any] = {
        status_column.name: status_value,
        "error_message": error[:1000] if error else None,
        "updated_at": int(time.time()),
    }
    if result is not None:
        values[result_column.name] = result
    if kind == "summary" and document_type:
        values["document_type"] = document_type
        values["document_type_source"] = document_type_source
    with _engine.begin() as conn:
        conn.execute(
            update(documents)
            .where(
                documents.c.id == document_id,
                documents.c.user_id == user_id,
            )
            .values(**values)
        )
        _set_task(conn, document_id, user_id, kind, status_value, error)


def _extract_pdf(path: Path) -> str:
    import fitz

    with fitz.open(str(path)) as document:
        return "\n\n".join(page.get_text("text") for page in document)


def _extract_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts = [
        paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_document(user_id: str, document_id: str) -> dict[str, Any]:
    row, cached = _acquire(document_id, user_id, "parse")
    if cached:
        return {
            "document_id": document_id,
            "text_length": len(row["extracted_text"]),
            "cached": True,
        }
    try:
        path = Path(row["storage_path"])
        text = _extract_pdf(path) if row["file_type"] == "pdf" else _extract_docx(path)
        text = text.strip()
        if not text:
            raise EmptyDocumentTextError("可能为扫描件，请使用 OCR")
        _finish(document_id, user_id, "parse", result=text)
        return {"document_id": document_id, "text_length": len(text), "cached": False}
    except EmptyDocumentTextError as exc:
        _finish(document_id, user_id, "parse", error=str(exc))
        raise
    except Exception as exc:
        message = f"文档解析失败：{exc}"
        _finish(document_id, user_id, "parse", error=message)
        raise DocumentSummaryError(message) from exc


def _normalize_list(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    return [str(item).strip() for item in value if str(item).strip()]


def _normalize_sections(value: Any) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []
    sections = []
    used_keys: set[str] = set()
    for index, item in enumerate(value[:6]):
        if not isinstance(item, dict):
            continue
        title = str(item.get("title") or "").strip()
        items = _normalize_list(item.get("items"))
        content = str(item.get("content") or "").strip()
        if content and not items:
            items = [content]
        if not title or not items:
            continue
        raw_key = str(item.get("key") or f"section_{index + 1}").strip().lower()
        key = "".join(
            character if character.isalnum() else "_" for character in raw_key
        )
        key = key.strip("_") or f"section_{index + 1}"
        if key in used_keys:
            key = f"{key}_{index + 1}"
        used_keys.add(key)
        sections.append({"key": key, "title": title, "items": items})
    return sections


def _legacy_sections(summary: dict[str, Any]) -> list[dict[str, Any]]:
    candidates = (
        ("key_points", "关键要点"),
        ("people", "相关人物"),
        ("dates", "日期与节点"),
        ("amounts", "金额与数据"),
        ("risks", "风险与待办"),
    )
    return [
        {"key": key, "title": title, "items": summary[key]}
        for key, title in candidates
        if summary[key]
    ]


def _normalize_summary(
    payload: Any, requested_document_type: str | None = None
) -> dict[str, Any]:
    if not isinstance(payload, dict) or not str(payload.get("summary") or "").strip():
        raise DocumentSummaryError("AI 返回的总结格式无效")
    detected_type = str(payload.get("document_type") or "").strip().lower()
    document_type = requested_document_type or (
        detected_type if detected_type in DOCUMENT_TYPES else "general"
    )
    result = {
        "summary": str(payload["summary"]).strip(),
        "key_points": _normalize_list(payload.get("key_points")),
        "people": _normalize_list(payload.get("people")),
        "dates": _normalize_list(payload.get("dates")),
        "amounts": _normalize_list(payload.get("amounts")),
        "risks": _normalize_list(payload.get("risks")),
        "document_type": document_type,
        "document_type_label": DOCUMENT_TYPE_LABELS[document_type],
        "risk_notice": RISK_NOTICES.get(document_type),
    }
    sections = _normalize_sections(payload.get("summary_sections"))
    if not sections:
        sections = _legacy_sections(result)
    result["summary_sections"] = sections[:6]
    if "source_truncated" in payload:
        result["source_truncated"] = bool(payload["source_truncated"])
    return result


async def call_deepseek(
    text: str, user_id: str, document_type: str | None = None
) -> dict[str, Any]:
    api_key = os.getenv("DEEPSEEK_API_KEY", "").strip()
    if not api_key:
        raise DocumentSummaryError("DeepSeek API 未配置")
    if len(text) > DEEPSEEK_MAX_INPUT_CHARS:
        head_length = int(DEEPSEEK_MAX_INPUT_CHARS * 0.7)
        tail_length = DEEPSEEK_MAX_INPUT_CHARS - head_length
        source_text = (
            text[:head_length]
            + "\n\n[文档中间部分因模型上下文限制已省略]\n\n"
            + text[-tail_length:]
        )
    else:
        source_text = text
    if document_type:
        type_instruction = (
            f"用户已指定文档类型为 {document_type}（{DOCUMENT_TYPE_LABELS[document_type]}），"
            "必须原样返回该 document_type，不要重新分类。"
        )
        section_guide = SUMMARY_SECTION_TEMPLATES[document_type]
    else:
        type_instruction = (
            "请从以下枚举中自动识别 document_type；无法可靠判断时返回 general："
            + ", ".join(DOCUMENT_TYPES)
            + "。"
        )
        section_guide = SUMMARY_SECTION_TEMPLATES
    prompt = (
        "请阅读下面的文档文本并输出严格 JSON，不得使用 Markdown 代码块。"
        f"{type_instruction}"
        "JSON 必须保留兼容字段：summary 字符串，以及 key_points、people、dates、amounts、risks 五个字符串数组。"
        "同时返回 summary_sections 数组，每项格式为 "
        '{"key":"snake_case","title":"卡片标题","items":["内容"]}。'
        "根据识别或指定类型，从最相关的维度生成 4 到 6 个卡片，卡片不得重复；"
        "没有依据的字段不要输出、不要臆测，空信息不要创建卡片。"
        "医疗、法律、财务文档必须在 risks 和对应风险卡片中明确列出限制、不确定性与需专业复核事项。"
        f"卡片维度参考：{json.dumps(section_guide, ensure_ascii=False, separators=(',', ':'))}\n\n"
        f"文档文本：\n{source_text}"
    )
    request_body = {
        "model": DEEPSEEK_MODEL,
        "messages": [
            {
                "role": "system",
                "content": "你是严谨的中文文档分析助手，只输出有效 JSON。",
            },
            {"role": "user", "content": prompt},
        ],
        "thinking": {"type": "disabled"},
        "response_format": {"type": "json_object"},
        "temperature": 0.2,
        "max_tokens": 3000,
        "stream": False,
        "user_id": user_id,
    }
    try:
        async with httpx.AsyncClient(timeout=DEEPSEEK_TIMEOUT_SECONDS) as client:
            response = await client.post(
                DEEPSEEK_API_URL,
                headers={"Authorization": f"Bearer {api_key}"},
                json=request_body,
            )
        response.raise_for_status()
        body = response.json()
        content = body["choices"][0]["message"]["content"]
        return _normalize_summary(json.loads(content), document_type)
    except httpx.TimeoutException as exc:
        raise DocumentSummaryError("AI 总结超时，请重试") from exc
    except httpx.HTTPStatusError as exc:
        raise DocumentSummaryError(
            f"DeepSeek API 请求失败（{exc.response.status_code}）"
        ) from exc
    except httpx.RequestError as exc:
        raise DocumentSummaryError("DeepSeek API 网络连接失败，请重试") from exc
    except (KeyError, IndexError, TypeError, json.JSONDecodeError) as exc:
        raise DocumentSummaryError("AI 返回的总结格式无效") from exc


async def summarize_document(
    user_id: str,
    document_id: str,
    *,
    document_type: str | None = None,
    regenerate: bool = False,
) -> dict[str, Any]:
    if document_type is not None:
        document_type = document_type.strip().lower()
        if document_type not in DOCUMENT_TYPES:
            raise DocumentSummaryError("不支持的文档类型")
    force = regenerate or document_type is not None
    row, cached = _acquire(document_id, user_id, "summary", force=force)
    if cached:
        stored_type = row.get("document_type")
        if stored_type not in DOCUMENT_TYPES:
            stored_type = None
        stored_summary = _normalize_summary(
            json.loads(row["summary_json"]), stored_type
        )
        return {"document_id": document_id, **stored_summary, "cached": True}
    current = get_owned_document(user_id, document_id)
    if current["extraction_status"] != "completed" or not current["extracted_text"]:
        error = "请先完成文档解析"
        _finish(document_id, user_id, "summary", error=error)
        raise DocumentSummaryError(error)
    try:
        summary = await call_deepseek(current["extracted_text"], user_id, document_type)
        # Keep the service boundary defensive: production DeepSeek responses and
        # test/custom adapters both pass through the same compatibility normalizer.
        summary = _normalize_summary(summary, document_type)
        summary["source_truncated"] = (
            len(current["extracted_text"]) > DEEPSEEK_MAX_INPUT_CHARS
        )
        serialized = json.dumps(summary, ensure_ascii=False, separators=(",", ":"))
        _finish(
            document_id,
            user_id,
            "summary",
            result=serialized,
            document_type=summary["document_type"],
            document_type_source="manual" if document_type else "auto",
        )
        return {"document_id": document_id, **summary, "cached": False}
    except DocumentSummaryError as exc:
        _finish(document_id, user_id, "summary", error=str(exc))
        raise
    except Exception as exc:
        error = "AI 总结失败，请重试"
        _finish(document_id, user_id, "summary", error=error)
        raise DocumentSummaryError(error) from exc


def save_to_history(user_id: str, document_id: str) -> None:
    get_owned_document(user_id, document_id)
    now = int(time.time())
    with _engine.begin() as conn:
        conn.execute(
            update(documents)
            .where(
                documents.c.id == document_id,
                documents.c.user_id == user_id,
            )
            .values(saved_at=now, updated_at=now)
        )


def list_history(user_id: str, limit: int = 30) -> list[dict[str, Any]]:
    init_user_database()
    with _engine.connect() as conn:
        rows = (
            conn.execute(
                select(documents)
                .where(
                    documents.c.user_id == user_id,
                    documents.c.saved_at.is_not(None),
                )
                .order_by(documents.c.saved_at.desc())
                .limit(limit)
            )
            .mappings()
            .all()
        )
    result = []
    for row in rows:
        summary = json.loads(row["summary_json"]) if row["summary_json"] else None
        result.append(
            {
                "document_id": row["id"],
                "filename": row["filename"],
                "file_type": row["file_type"],
                "file_size": row["file_size"],
                "summary": summary,
                "saved_at": row["saved_at"],
            }
        )
    return result
