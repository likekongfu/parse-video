"""Cached paragraph-oriented document translation backed by DeepSeek."""

from __future__ import annotations

import asyncio
import hashlib
import html
import io
import json
import logging
import os
import re
import time
import uuid
from pathlib import Path
from typing import Any, Callable, Iterable

import httpx
from docx import Document
from docx.oxml.ns import qn
from sqlalchemy import and_, delete, insert, or_, select, update
from sqlalchemy.exc import IntegrityError

from parse_video_py.document_summary import get_owned_document
from parse_video_py.user_db import (
    _engine,
    document_translation_jobs,
    document_translations,
    init_user_database,
)

PROCESSING_STALE_SECONDS = 10 * 60
TRANSLATION_BATCH_SIZE = max(1, int(os.getenv("TRANSLATION_BATCH_SIZE", "20")))
TRANSLATION_BATCH_RETRIES = 2
TRANSLATION_PIPELINE_VERSION = "structure-location-v3"
TRANSLATION_JOB_TIMEOUT_SECONDS = max(
    30, int(os.getenv("TRANSLATION_JOB_TIMEOUT_SECONDS", "1800"))
)
TRANSLATION_JOB_TTL_SECONDS = max(
    60, int(os.getenv("TRANSLATION_JOB_TTL_SECONDS", "7200"))
)
TRANSLATION_JOB_CLEANUP_INTERVAL_SECONDS = max(
    10, int(os.getenv("TRANSLATION_JOB_CLEANUP_INTERVAL_SECONDS", "60"))
)
PDF_TRANSLATION_MIN_SCALE = min(
    1.0, max(0.25, float(os.getenv("PDF_TRANSLATION_MIN_SCALE", "0.35")))
)
PDF_TRANSLATION_ABSOLUTE_MIN_SCALE = 0.25
PDF_LAYOUT_GAP = 2.0
PDF_LAYOUT_HORIZONTAL_EXPANSION = 18.0
DEEPSEEK_API_URL = os.getenv(
    "DEEPSEEK_API_URL", "https://api.deepseek.com/chat/completions"
).strip()
DEEPSEEK_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-v4-flash").strip()
DEEPSEEK_TIMEOUT_SECONDS = float(os.getenv("DEEPSEEK_TIMEOUT_SECONDS", "120"))

ALLOWED_MODES = {"translation", "bilingual"}
ALLOWED_STYLES = {"general", "business", "legal", "technical"}
STYLE_LABELS = {
    "general": "通用：自然、准确、清晰",
    "business": "商务：专业、简洁、正式",
    "legal": "法律：严谨、保守，不弱化义务、限制和条件",
    "technical": "技术：准确保留技术术语、参数、单位和代码",
}
_LANGUAGE_CODE = re.compile(r"^[A-Za-z][A-Za-z0-9-]{0,15}$")
logger = logging.getLogger(__name__)
_background_jobs: set[asyncio.Task[None]] = set()
_cleanup_task: asyncio.Task[None] | None = None


class DocumentTranslationError(RuntimeError):
    pass


class DocumentTranslationBusyError(DocumentTranslationError):
    pass


class InvalidTranslationRequestError(DocumentTranslationError):
    """The requested mode or export format is incompatible with the source file."""


class PDFLayoutOverflowError(DocumentTranslationError):
    """Translated text cannot fit safely inside the original PDF layout."""

    error_code = "PDF_LAYOUT_OVERFLOW"

    def __init__(
        self,
        *,
        page_number: int,
        segment_id: str,
        original_bbox: list[float] | tuple[float, ...],
        attempted_bbox: list[float] | tuple[float, ...],
        attempted_scale: float,
        translated_length: int,
    ) -> None:
        self.page_number = page_number
        self.segment_id = segment_id
        self.original_bbox = [round(float(value), 3) for value in original_bbox]
        self.attempted_bbox = [round(float(value), 3) for value in attempted_bbox]
        self.attempted_scale = round(float(attempted_scale), 3)
        self.translated_length = translated_length
        super().__init__(
            "译文过长，无法完整放入原版式区域 "
            f"(error_code={self.error_code}, page_number={page_number}, "
            f"segment_id={segment_id})"
        )

    def to_detail(self) -> dict[str, Any]:
        return {
            "error_code": self.error_code,
            "message": "译文过长，无法完整放入原版式区域",
            "page_number": self.page_number,
            "segment_id": self.segment_id,
            "original_bbox": self.original_bbox,
            "attempted_bbox": self.attempted_bbox,
            "attempted_scale": self.attempted_scale,
            "translated_length": self.translated_length,
        }


class InvalidTranslationResponseError(DocumentTranslationError):
    """The model response cannot be decoded as a translation payload."""

    def __init__(
        self,
        message: str,
        *,
        finish_reason: str | None = None,
        expected_count: int | None = None,
        actual_count: int | None = None,
    ) -> None:
        super().__init__(message)
        self.finish_reason = finish_reason
        self.expected_count = expected_count
        self.actual_count = actual_count


class TruncatedTranslationResponseError(InvalidTranslationResponseError):
    """The model stopped because its output token limit was reached."""


class IncompleteTranslationResponseError(DocumentTranslationError):
    """The model returned only a usable subset of the requested segments."""

    def __init__(
        self,
        translated: list[dict[str, str]],
        detected_source_language: str | None,
        *,
        finish_reason: str | None = None,
        expected_count: int | None = None,
        actual_count: int | None = None,
        missing_ids: list[str] | None = None,
        extra_ids: list[str] | None = None,
        duplicate_ids: list[str] | None = None,
    ) -> None:
        super().__init__("AI 返回的翻译段落不完整")
        self.translated = translated
        self.detected_source_language = detected_source_language
        self.finish_reason = finish_reason
        self.expected_count = expected_count
        self.actual_count = actual_count
        self.missing_ids = missing_ids or []
        self.extra_ids = extra_ids or []
        self.duplicate_ids = duplicate_ids or []


def normalize_glossary(value: Any) -> list[dict[str, str]]:
    if value in (None, ""):
        return []
    if not isinstance(value, list):
        raise DocumentTranslationError("术语表格式无效")
    result: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()
    for item in value:
        if not isinstance(item, dict):
            raise DocumentTranslationError("术语表格式无效")
        source = str(item.get("source") or "").strip()
        target = str(item.get("target") or "").strip()
        if not source or not target:
            raise DocumentTranslationError("术语表的原词和译词不能为空")
        if len(source) > 120 or len(target) > 240:
            raise DocumentTranslationError("单条术语过长")
        key = (source, target)
        if key not in seen:
            result.append({"source": source, "target": target})
            seen.add(key)
    if len(result) > 100:
        raise DocumentTranslationError("术语表最多支持 100 条")
    return result


def normalize_options(
    *,
    source_language: str,
    target_language: str,
    mode: str,
    style: str,
    glossary: Any,
) -> dict[str, Any]:
    source_language = (source_language or "auto").strip()
    target_language = (target_language or "").strip()
    mode = (mode or "translation").strip()
    style = (style or "general").strip()
    if source_language != "auto" and not _LANGUAGE_CODE.fullmatch(source_language):
        raise DocumentTranslationError("源语言无效")
    if not _LANGUAGE_CODE.fullmatch(target_language):
        raise DocumentTranslationError("请选择有效的目标语言")
    if mode not in ALLOWED_MODES:
        raise DocumentTranslationError("翻译模式无效")
    if style not in ALLOWED_STYLES:
        raise DocumentTranslationError("翻译风格无效")
    return {
        "source_language": source_language,
        "target_language": target_language,
        "mode": mode,
        "style": style,
        "glossary": normalize_glossary(glossary),
    }


def _options_hash(options: dict[str, Any]) -> bytes:
    canonical = dict(options)
    canonical["pipeline_version"] = TRANSLATION_PIPELINE_VERSION
    canonical["glossary"] = sorted(
        options["glossary"], key=lambda item: (item["source"], item["target"])
    )
    serialized = json.dumps(
        canonical, ensure_ascii=False, sort_keys=True, separators=(",", ":")
    )
    return hashlib.sha256(serialized.encode("utf-8")).digest()


def _segment_kind(text: str) -> str:
    compact = re.sub(r"\s+", " ", text).strip()
    is_heading = len(compact) <= 80 and (
        bool(
            re.match(
                r"^(第[一二三四五六七八九十百0-9]+[章节条]|[0-9一二三四五六七八九十]+[.、])",
                compact,
            )
        )
        or not re.search(r"[。！？；.!?;]$", compact)
    )
    return "heading" if is_heading else "paragraph"


def segment_document(text: str) -> list[dict[str, str]]:
    """Split on document line/paragraph boundaries, never on sentences."""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    blocks = [block.strip() for block in re.split(r"\n+", normalized) if block.strip()]
    result = []
    for index, block in enumerate(blocks, 1):
        result.append(
            {
                "segment_id": f"seg-{index:04d}",
                "source_text": block,
                "kind": _segment_kind(block),
            }
        )
    if not result:
        raise DocumentTranslationError("可能为扫描件，请使用 OCR")
    return result


_DOCX_EXCLUDED_TEXT_ANCESTORS = {
    qn("w:drawing"),
    qn("w:object"),
    qn("w:pict"),
    qn("w:fldSimple"),
}


def _docx_node_is_excluded(node) -> bool:
    return any(
        ancestor.tag in _DOCX_EXCLUDED_TEXT_ANCESTORS
        for ancestor in node.iterancestors()
    )


def _docx_paragraph_tokens(paragraph) -> list[tuple[str, Any]]:
    """Return visible text and structural separators without touching field/drawing XML."""
    tokens: list[tuple[str, Any]] = []
    field_depth = 0
    for node in paragraph._p.iter():
        if node.tag == qn("w:fldChar"):
            field_type = node.get(qn("w:fldCharType"))
            if field_type == "begin":
                field_depth += 1
            elif field_type == "end":
                field_depth = max(0, field_depth - 1)
            continue
        if field_depth or _docx_node_is_excluded(node):
            continue
        if node.tag == qn("w:t"):
            tokens.append(("text", node))
        elif node.tag == qn("w:tab"):
            tokens.append(("tab", node))
        elif node.tag in (qn("w:br"), qn("w:cr")):
            tokens.append(("break", node))
    return tokens


def _docx_paragraph_text(paragraph) -> str:
    parts: list[str] = []
    for token_type, node in _docx_paragraph_tokens(paragraph):
        if token_type == "text":
            parts.append(node.text or "")
        elif token_type == "tab":
            parts.append("\t")
        else:
            parts.append("\n")
    return "".join(parts)


def _docx_location_id(location: dict[str, Any]) -> str:
    scope = location["scope"]
    if scope == "body":
        return f"body-p{location['paragraph_index'] + 1:04d}"
    if scope == "table":
        return (
            f"table-t{location['table_index'] + 1:04d}"
            f"-r{location['row_index'] + 1:04d}"
            f"-c{location['cell_index'] + 1:04d}"
            f"-p{location['paragraph_index'] + 1:04d}"
        )
    prefix = f"{scope}-s{location['section_index'] + 1:04d}" f"-{location['variant']}"
    if scope in ("header", "footer"):
        return f"{prefix}-p{location['paragraph_index'] + 1:04d}"
    return (
        f"{prefix}-t{location['table_index'] + 1:04d}"
        f"-r{location['row_index'] + 1:04d}"
        f"-c{location['cell_index'] + 1:04d}"
        f"-p{location['paragraph_index'] + 1:04d}"
    )


def _append_docx_segment(
    segments: list[dict[str, Any]],
    paragraph,
    location: dict[str, Any],
    seen_paragraphs: set[Any],
) -> None:
    paragraph_element = paragraph._p
    if paragraph_element in seen_paragraphs:
        return
    seen_paragraphs.add(paragraph_element)
    source_text = _docx_paragraph_text(paragraph)
    if not source_text.strip():
        return
    segments.append(
        {
            "segment_id": _docx_location_id(location),
            "source_text": source_text,
            "kind": _segment_kind(source_text),
            "location": location,
        }
    )


def _append_docx_table_segments(
    segments: list[dict[str, Any]],
    tables,
    seen_paragraphs: set[Any],
    *,
    scope: str,
    section_index: int | None = None,
    variant: str | None = None,
) -> None:
    for table_index, table in enumerate(tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    location: dict[str, Any] = {
                        "scope": scope,
                        "table_index": table_index,
                        "row_index": row_index,
                        "cell_index": cell_index,
                        "paragraph_index": paragraph_index,
                    }
                    if section_index is not None:
                        location["section_index"] = section_index
                    if variant is not None:
                        location["variant"] = variant
                    _append_docx_segment(segments, paragraph, location, seen_paragraphs)


def _docx_story_references(section, story: str):
    reference_tag = qn(f"w:{story}Reference")
    references = [child for child in section._sectPr if child.tag == reference_tag]
    for reference in references:
        variant = reference.get(qn("w:type"), "default")
        if story == "header":
            container = (
                section.first_page_header
                if variant == "first"
                else section.even_page_header if variant == "even" else section.header
            )
        else:
            container = (
                section.first_page_footer
                if variant == "first"
                else section.even_page_footer if variant == "even" else section.footer
            )
        yield variant, container


def segment_docx_document(path: Path) -> list[dict[str, Any]]:
    """Create translation segments from stable DOCX structure locations."""
    source_document = Document(str(path))
    segments: list[dict[str, Any]] = []
    seen_paragraphs: set[Any] = set()
    for paragraph_index, paragraph in enumerate(source_document.paragraphs):
        _append_docx_segment(
            segments,
            paragraph,
            {"scope": "body", "paragraph_index": paragraph_index},
            seen_paragraphs,
        )
    _append_docx_table_segments(
        segments,
        source_document.tables,
        seen_paragraphs,
        scope="table",
    )

    seen_story_parts: set[str] = set()
    for section_index, section in enumerate(source_document.sections):
        for story in ("header", "footer"):
            for variant, container in _docx_story_references(section, story):
                part_key = str(container.part.partname)
                if part_key in seen_story_parts:
                    continue
                seen_story_parts.add(part_key)
                for paragraph_index, paragraph in enumerate(container.paragraphs):
                    _append_docx_segment(
                        segments,
                        paragraph,
                        {
                            "scope": story,
                            "section_index": section_index,
                            "variant": variant,
                            "paragraph_index": paragraph_index,
                        },
                        seen_paragraphs,
                    )
                _append_docx_table_segments(
                    segments,
                    container.tables,
                    seen_paragraphs,
                    scope=f"{story}_table",
                    section_index=section_index,
                    variant=variant,
                )
    if not segments:
        raise DocumentTranslationError("文档中没有可翻译文字")
    return segments


def _pdf_block_alignment(block_bbox: tuple[float, ...], lines: list[dict]) -> str:
    block_left, _, block_right, _ = block_bbox
    block_width = max(1.0, block_right - block_left)
    line_boxes = [line.get("bbox") for line in lines if line.get("bbox")]
    if not line_boxes:
        return "left"
    left_gap = min(float(box[0]) for box in line_boxes) - block_left
    right_gap = block_right - max(float(box[2]) for box in line_boxes)
    if abs(left_gap - right_gap) <= max(4.0, block_width * 0.08):
        return "center"
    if left_gap > right_gap * 2 + 4:
        return "right"
    return "left"


def segment_pdf_document(path: Path) -> list[dict[str, Any]]:
    """Extract stable page/block segments and the geometry needed for PDF export."""
    import fitz

    segments: list[dict[str, Any]] = []
    with fitz.open(str(path)) as source:
        for page_index, page in enumerate(source):
            page_dict = page.get_text("dict", sort=True)
            text_block_index = 0
            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:
                    continue
                lines = block.get("lines") or []
                line_texts: list[str] = []
                spans: list[dict[str, Any]] = []
                for line in lines:
                    line_spans = line.get("spans") or []
                    spans.extend(line_spans)
                    line_text = "".join(
                        str(span.get("text") or "") for span in line_spans
                    )
                    if line_text.strip():
                        line_texts.append(line_text.strip())
                source_text = "\n".join(line_texts).strip()
                if not source_text:
                    continue
                bbox = tuple(float(value) for value in block.get("bbox", ()))
                if len(bbox) != 4 or bbox[2] - bbox[0] < 1 or bbox[3] - bbox[1] < 1:
                    continue
                text_block_index += 1
                primary_span = max(
                    spans,
                    key=lambda span: len(str(span.get("text") or "")),
                    default={},
                )
                font_size = max(
                    4.0,
                    float(primary_span.get("size") or max(8.0, bbox[3] - bbox[1])),
                )
                color = int(primary_span.get("color") or 0) & 0xFFFFFF
                font_name = str(primary_span.get("font") or "")
                segments.append(
                    {
                        "segment_id": f"p{page_index + 1:04d}-b{text_block_index:04d}",
                        "source_text": source_text,
                        "kind": _segment_kind(source_text),
                        "layout": {
                            "page_number": page_index + 1,
                            "block_index": text_block_index,
                            "bbox": [round(value, 3) for value in bbox],
                            "font_size": round(font_size, 2),
                            "font_color": f"#{color:06x}",
                            "font_family": (
                                "monospace"
                                if "courier" in font_name.lower()
                                else (
                                    "serif"
                                    if any(
                                        name in font_name.lower()
                                        for name in ("times", "serif", "song", "ming")
                                    )
                                    else "sans-serif"
                                )
                            ),
                            "text_align": _pdf_block_alignment(bbox, lines),
                        },
                    }
                )
    if not segments:
        raise DocumentTranslationError("可能为扫描件，请使用 OCR")
    return segments


def iter_segment_batches(
    segments: list[dict[str, str]], batch_size: int = TRANSLATION_BATCH_SIZE
) -> Iterable[list[dict[str, str]]]:
    """Yield ordered batches capped by segment count, never by output size guesses."""
    if batch_size < 1:
        raise ValueError("batch_size must be at least 1")
    for offset in range(0, len(segments), batch_size):
        yield segments[offset : offset + batch_size]


_URL_ONLY = re.compile(r"^(?:https?://|www\.)\S+$", re.IGNORECASE)
_EMAIL_ONLY = re.compile(r"^[^\s@]+@[^\s@]+\.[^\s@]+$")
_NUMERIC_ONLY = re.compile(r"^[\d\s.,，。:：;；+\-−–—/%％‰()（）\[\]{}￥¥$€£]+$")


def segment_requires_translation(segment: dict[str, Any]) -> bool:
    """Return False for content that must be preserved without calling the model."""
    text = str(segment.get("source_text") or "").strip()
    if not text:
        return False
    if (
        _URL_ONLY.fullmatch(text)
        or _EMAIL_ONLY.fullmatch(text)
        or _NUMERIC_ONLY.fullmatch(text)
    ):
        return False
    # isalpha() covers letters across supported scripts without classifying digits
    # and underscore as translatable content.
    return any(character.isalpha() for character in text)


def _passthrough_translation(segment: dict[str, Any]) -> dict[str, str]:
    return {
        "segment_id": str(segment["segment_id"]),
        "translated_text": str(segment.get("source_text") or ""),
    }


def _normalize_ai_result(
    payload: Any, expected_segments: list[dict[str, str]]
) -> tuple[str | None, list[dict[str, str]]]:
    expected_count = len(expected_segments)
    if not isinstance(payload, dict) or not isinstance(
        payload.get("translations"), list
    ):
        raise InvalidTranslationResponseError(
            "AI 返回的翻译格式无效",
            expected_count=expected_count,
            actual_count=None,
        )
    expected_ids = [item["segment_id"] for item in expected_segments]
    expected_id_set = set(expected_ids)
    response_items = payload["translations"]
    actual_count = len(response_items)
    translated: dict[str, str] = {}
    response_ids: list[str] = []
    invalid_item = False
    for item in response_items:
        if not isinstance(item, dict):
            invalid_item = True
            continue
        segment_id = str(item.get("segment_id") or "")
        text = str(item.get("translated_text") or "").strip()
        response_ids.append(segment_id)
        if segment_id in translated or segment_id not in expected_id_set or not text:
            continue
        translated[segment_id] = text
    detected = str(payload.get("detected_source_language") or "").strip() or None
    id_counts = {
        segment_id: response_ids.count(segment_id) for segment_id in response_ids
    }
    duplicate_ids = sorted(
        segment_id
        for segment_id, count in id_counts.items()
        if segment_id and count > 1
    )
    extra_ids = sorted(
        {segment_id for segment_id in response_ids if segment_id not in expected_id_set}
    )
    missing_ids = [
        segment_id for segment_id in expected_ids if segment_id not in translated
    ]
    if (
        invalid_item
        or missing_ids
        or extra_ids
        or duplicate_ids
        or actual_count != expected_count
    ):
        partial = [
            {"segment_id": segment_id, "translated_text": translated[segment_id]}
            for segment_id in expected_ids
            if segment_id in translated
        ]
        raise IncompleteTranslationResponseError(
            partial,
            detected,
            expected_count=expected_count,
            actual_count=actual_count,
            missing_ids=missing_ids,
            extra_ids=extra_ids,
            duplicate_ids=duplicate_ids,
        )
    return detected, [
        {"segment_id": segment_id, "translated_text": translated[segment_id]}
        for segment_id in expected_ids
    ]


async def call_deepseek_translation(
    segments: list[dict[str, str]],
    *,
    source_language: str,
    target_language: str,
    style: str,
    glossary: list[dict[str, str]],
    user_id: str,
    request_id: str = "unknown",
    batch_index: int = 0,
    attempt: int = 0,
) -> tuple[str | None, list[dict[str, str]]]:
    api_key = os.getenv("DEEPSEEK_API_KEY", "").strip()
    if not api_key:
        raise DocumentTranslationError("DeepSeek API 未配置")
    source_instruction = (
        "自动识别源语言"
        if source_language == "auto"
        else f"源语言代码为 {source_language}"
    )
    glossary_text = json.dumps(glossary, ensure_ascii=False, separators=(",", ":"))
    source_segments = [
        {
            "segment_id": item["segment_id"],
            "kind": item["kind"],
            "text": item["source_text"],
        }
        for item in segments
    ]
    required_ids = [item["segment_id"] for item in segments]
    prompt = (
        f"将以下文档段落翻译为 {target_language}。{source_instruction}。"
        f"风格要求：{STYLE_LABELS[style]}。一次处理的是完整标题和段落，禁止拆成逐句结果。"
        f"输入共有 {len(required_ids)} 个 segment。必须为每个输入 segment_id 返回且只返回一项。"
        "禁止遗漏、合并、改名、重复、新增或重排 segment_id；即使译文与原文相同也必须返回。"
        f"必须严格按以下顺序返回 segment_id："
        f"{json.dumps(required_ids, ensure_ascii=False, separators=(',', ':'))}。"
        "原样保留编号、金额、货币、日期、计量单位、网址、邮箱、代码和无法确定的专有名词。"
        "优先严格使用术语表。只能输出一个合法 JSON 对象，禁止 Markdown、代码块、解释或说明文字。"
        "顶层只能包含 translations；每个返回项只能包含 segment_id 和 translated_text。"
        'JSON 格式为 {"translations":[{"segment_id":"seg-0001",'
        '"translated_text":"..."}]}。\n'
        f"术语表：{glossary_text}\n"
        f"段落：{json.dumps(source_segments, ensure_ascii=False, separators=(',', ':'))}"
    )
    request_body = {
        "model": DEEPSEEK_MODEL,
        "messages": [
            {
                "role": "system",
                "content": (
                    "你是专业文档翻译引擎。输出必须是严格有效的 JSON，"
                    "不得输出 Markdown、代码围栏、自然语言说明或任何 JSON 之外的字符。"
                ),
            },
            {"role": "user", "content": prompt},
        ],
        "thinking": {"type": "disabled"},
        "response_format": {"type": "json_object"},
        "temperature": 0.1,
        "max_tokens": 8000,
        "stream": False,
        "user_id": user_id,
    }
    try:
        async with httpx.AsyncClient(timeout=DEEPSEEK_TIMEOUT_SECONDS) as client:
            response = await client.post(
                DEEPSEEK_API_URL,
                headers={"Authorization": f"Bearer {api_key}"},
                json=request_body,
            )
        response.raise_for_status()
        choice = response.json()["choices"][0]
        finish_reason = str(choice.get("finish_reason") or "unknown")
        logger.info(
            "translation_batch_response request_id=%s batch_index=%d "
            "segment_count=%d attempt=%d finish_reason=%s",
            request_id,
            batch_index,
            len(segments),
            attempt,
            finish_reason,
        )
        if finish_reason == "length":
            raise TruncatedTranslationResponseError(
                "AI 翻译输出被截断", finish_reason=finish_reason
            )
        content = choice["message"]["content"]
        try:
            payload = json.loads(content)
            return _normalize_ai_result(payload, segments)
        except IncompleteTranslationResponseError as exc:
            exc.finish_reason = finish_reason
            logger.warning(
                "translation_response_validation_failed request_id=%s batch_index=%d "
                "attempt=%d finish_reason=%s expected_count=%d actual_count=%d "
                "missing_ids=%s extra_ids=%s duplicate_ids=%s",
                request_id,
                batch_index,
                attempt,
                finish_reason,
                exc.expected_count,
                exc.actual_count,
                json.dumps(exc.missing_ids, ensure_ascii=True),
                json.dumps(exc.extra_ids, ensure_ascii=True),
                json.dumps(exc.duplicate_ids, ensure_ascii=True),
            )
            raise
        except InvalidTranslationResponseError as exc:
            exc.finish_reason = finish_reason
            logger.warning(
                "translation_response_validation_failed request_id=%s batch_index=%d "
                "attempt=%d finish_reason=%s expected_count=%s actual_count=%s "
                "missing_ids=[] extra_ids=[] duplicate_ids=[] reason=invalid_schema",
                request_id,
                batch_index,
                attempt,
                finish_reason,
                exc.expected_count,
                exc.actual_count,
            )
            raise
        except json.JSONDecodeError as exc:
            logger.warning(
                "translation_response_validation_failed request_id=%s batch_index=%d "
                "attempt=%d finish_reason=%s expected_count=%d actual_count=unknown "
                "missing_ids=%s extra_ids=[] duplicate_ids=[] reason=invalid_json "
                "json_error_line=%d json_error_column=%d json_error_position=%d",
                request_id,
                batch_index,
                attempt,
                finish_reason,
                len(segments),
                json.dumps(required_ids, ensure_ascii=True),
                exc.lineno,
                exc.colno,
                exc.pos,
            )
            raise InvalidTranslationResponseError(
                "AI 返回的翻译 JSON 无效",
                finish_reason=finish_reason,
                expected_count=len(segments),
                actual_count=None,
            ) from exc
    except httpx.TimeoutException as exc:
        logger.warning(
            "translation_batch_request_failed request_id=%s batch_index=%d "
            "segment_count=%d attempt=%d finish_reason=unknown reason=timeout",
            request_id,
            batch_index,
            len(segments),
            attempt,
        )
        raise DocumentTranslationError("AI 翻译超时，请重试") from exc
    except httpx.HTTPStatusError as exc:
        logger.warning(
            "translation_batch_request_failed request_id=%s batch_index=%d "
            "segment_count=%d attempt=%d finish_reason=unknown reason=http_%d",
            request_id,
            batch_index,
            len(segments),
            attempt,
            exc.response.status_code,
        )
        raise DocumentTranslationError(
            f"DeepSeek API 请求失败（{exc.response.status_code}）"
        ) from exc
    except httpx.RequestError as exc:
        logger.warning(
            "translation_batch_request_failed request_id=%s batch_index=%d "
            "segment_count=%d attempt=%d finish_reason=unknown reason=network",
            request_id,
            batch_index,
            len(segments),
            attempt,
        )
        raise DocumentTranslationError("DeepSeek API 网络连接失败，请重试") from exc
    except (KeyError, IndexError, TypeError) as exc:
        logger.warning(
            "translation_response_validation_failed request_id=%s batch_index=%d "
            "attempt=%d finish_reason=unknown expected_count=%d actual_count=unknown "
            "missing_ids=%s extra_ids=[] duplicate_ids=[] reason=invalid_schema",
            request_id,
            batch_index,
            attempt,
            len(segments),
            json.dumps(required_ids, ensure_ascii=True),
        )
        raise InvalidTranslationResponseError(
            "AI 返回的翻译格式无效",
            expected_count=len(segments),
            actual_count=None,
        ) from exc


async def translate_batch_with_recovery(
    segments: list[dict[str, str]],
    *,
    source_language: str,
    target_language: str,
    style: str,
    glossary: list[dict[str, str]],
    user_id: str,
    request_id: str = "unknown",
    batch_index: int = 0,
    max_retries: int = TRANSLATION_BATCH_RETRIES,
    split_depth: int = 0,
) -> tuple[str | None, list[dict[str, str]]]:
    """Retry a batch at most twice, then bisect it until a clear singleton error."""
    ordered_segments = [
        segment for segment in segments if str(segment.get("source_text") or "").strip()
    ]
    if not ordered_segments:
        return None, []
    passthrough_by_id = {
        str(segment["segment_id"]): _passthrough_translation(segment)
        for segment in ordered_segments
        if not segment_requires_translation(segment)
    }
    model_segments = [
        segment for segment in ordered_segments if segment_requires_translation(segment)
    ]
    logger.debug(
        "translation_batch_prepared request_id=%s batch_index=%d input_count=%d "
        "filtered_blank_count=%d passthrough_count=%d model_count=%d split_depth=%d",
        request_id,
        batch_index,
        len(segments),
        len(segments) - len(ordered_segments),
        len(passthrough_by_id),
        len(model_segments),
        split_depth,
    )
    if not model_segments:
        return None, [
            passthrough_by_id[str(segment["segment_id"])]
            for segment in ordered_segments
        ]

    retry_limit = min(TRANSLATION_BATCH_RETRIES, max(0, max_retries))
    last_error: (
        InvalidTranslationResponseError | IncompleteTranslationResponseError | None
    ) = None
    for attempt in range(retry_limit + 1):
        try:
            detected, translated = await call_deepseek_translation(
                model_segments,
                source_language=source_language,
                target_language=target_language,
                style=style,
                glossary=glossary,
                user_id=user_id,
                request_id=request_id,
                batch_index=batch_index,
                attempt=attempt,
            )
            translated_by_id = {
                item["segment_id"]: item for item in translated
            } | passthrough_by_id
            return detected, [
                translated_by_id[str(segment["segment_id"])]
                for segment in ordered_segments
            ]
        except (
            InvalidTranslationResponseError,
            IncompleteTranslationResponseError,
        ) as exc:
            last_error = exc
            logger.warning(
                "translation_batch_retry request_id=%s batch_index=%d "
                "segment_count=%d attempt=%d finish_reason=%s split_depth=%d reason=%s "
                "expected_count=%s actual_count=%s missing_ids=%s extra_ids=%s "
                "duplicate_ids=%s",
                request_id,
                batch_index,
                len(model_segments),
                attempt,
                getattr(exc, "finish_reason", None) or "unknown",
                split_depth,
                type(exc).__name__,
                getattr(exc, "expected_count", None),
                getattr(exc, "actual_count", None),
                json.dumps(getattr(exc, "missing_ids", []), ensure_ascii=True),
                json.dumps(getattr(exc, "extra_ids", []), ensure_ascii=True),
                json.dumps(getattr(exc, "duplicate_ids", []), ensure_ascii=True),
            )

    if len(model_segments) == 1:
        segment_id = model_segments[0]["segment_id"]
        logger.error(
            "translation_segment_failed request_id=%s batch_index=%d "
            "failed_segment_id=%s retry_count=%d split_depth=%d reason=%s",
            request_id,
            batch_index,
            segment_id,
            retry_limit,
            split_depth,
            type(last_error).__name__ if last_error else "unknown",
        )
        raise DocumentTranslationError(
            f"段落 {segment_id} 翻译失败：AI 返回内容无效或不完整"
        ) from last_error

    midpoint = len(model_segments) // 2
    logger.warning(
        "translation_batch_split request_id=%s batch_index=%d segment_count=%d "
        "finish_reason=%s split_depth=%d",
        request_id,
        batch_index,
        len(model_segments),
        getattr(last_error, "finish_reason", None) or "unknown",
        split_depth,
    )
    detected_language: str | None = None
    translated_by_id: dict[str, dict[str, str]] = dict(passthrough_by_id)
    for smaller_batch in (
        model_segments[:midpoint],
        model_segments[midpoint:],
    ):
        detected, translated = await translate_batch_with_recovery(
            smaller_batch,
            source_language=source_language,
            target_language=target_language,
            style=style,
            glossary=glossary,
            user_id=user_id,
            request_id=request_id,
            batch_index=batch_index,
            max_retries=max_retries,
            split_depth=split_depth + 1,
        )
        detected_language = detected_language or detected
        translated_by_id.update((item["segment_id"], item) for item in translated)
    return detected_language, [
        translated_by_id[str(segment["segment_id"])] for segment in ordered_segments
    ]


def _translation_row(user_id: str, document_id: str, options_hash: bytes):
    with _engine.connect() as conn:
        return (
            conn.execute(
                select(document_translations).where(
                    document_translations.c.user_id == user_id,
                    document_translations.c.document_id == document_id,
                    document_translations.c.options_hash == options_hash,
                )
            )
            .mappings()
            .first()
        )


def _acquire_translation(
    user_id: str, document_id: str, options: dict[str, Any]
) -> tuple[dict[str, Any], bool]:
    init_user_database()
    options_hash = _options_hash(options)
    now = int(time.time())
    row = _translation_row(user_id, document_id, options_hash)
    if row and row["status"] == "completed" and row["result_json"]:
        return dict(row), True
    if row:
        with _engine.begin() as conn:
            acquired = conn.execute(
                update(document_translations)
                .where(
                    document_translations.c.id == row["id"],
                    or_(
                        document_translations.c.status == "failed",
                        and_(
                            document_translations.c.status == "processing",
                            document_translations.c.updated_at
                            < now - PROCESSING_STALE_SECONDS,
                        ),
                    ),
                )
                .values(status="processing", error_message=None, updated_at=now)
            )
        if acquired.rowcount != 1:
            raise DocumentTranslationBusyError("文档正在翻译中，请稍后重试")
        return dict(row), False

    values = {
        "id": str(uuid.uuid4()),
        "document_id": document_id,
        "user_id": user_id,
        "options_hash": options_hash,
        "source_language": options["source_language"],
        "detected_source_language": None,
        "target_language": options["target_language"],
        "mode": options["mode"],
        "style": options["style"],
        "glossary_json": json.dumps(
            options["glossary"], ensure_ascii=False, separators=(",", ":")
        ),
        "result_json": None,
        "status": "processing",
        "error_message": None,
        "created_at": now,
        "updated_at": now,
        "completed_at": None,
    }
    try:
        with _engine.begin() as conn:
            conn.execute(insert(document_translations).values(**values))
        return values, False
    except IntegrityError:
        winner = _translation_row(user_id, document_id, options_hash)
        if winner and winner["status"] == "completed" and winner["result_json"]:
            return dict(winner), True
        raise DocumentTranslationBusyError("文档正在翻译中，请稍后重试")


def _finish_translation(
    translation_id: str,
    user_id: str,
    *,
    result: dict[str, Any] | None = None,
    detected_source_language: str | None = None,
    error: str | None = None,
) -> None:
    now = int(time.time())
    values: dict[str, Any] = {
        "status": "failed" if error else "completed",
        "error_message": error[:1000] if error else None,
        "updated_at": now,
        "completed_at": None if error else now,
    }
    if result is not None:
        values["result_json"] = json.dumps(
            result, ensure_ascii=False, separators=(",", ":")
        )
        values["detected_source_language"] = detected_source_language
    with _engine.begin() as conn:
        conn.execute(
            update(document_translations)
            .where(
                document_translations.c.id == translation_id,
                document_translations.c.user_id == user_id,
            )
            .values(**values)
        )


async def translate_document(
    user_id: str,
    document_id: str,
    *,
    source_language: str = "auto",
    target_language: str,
    mode: str = "translation",
    style: str = "general",
    glossary: Any = None,
    request_id: str | None = None,
    progress_callback: Callable[[int, int], None] | None = None,
) -> dict[str, Any]:
    document = get_owned_document(user_id, document_id)
    if document["extraction_status"] != "completed" or not document["extracted_text"]:
        raise DocumentTranslationError("请先完成文档解析")
    options = normalize_options(
        source_language=source_language,
        target_language=target_language,
        mode=mode,
        style=style,
        glossary=glossary,
    )
    if document.get("file_type") == "pdf" and options["mode"] != "translation":
        raise InvalidTranslationRequestError("PDF翻译暂不支持双语模式")
    task, cached = _acquire_translation(user_id, document_id, options)
    if cached:
        if progress_callback:
            progress_callback(1, 1)
        return {
            "document_id": document_id,
            "translation_id": task["id"],
            **json.loads(task["result_json"]),
            "cached": True,
        }
    request_id = request_id or uuid.uuid4().hex
    try:
        if document.get("file_type") == "pdf":
            segments = segment_pdf_document(Path(document["storage_path"]))
        elif document.get("file_type") == "docx":
            segments = segment_docx_document(Path(document["storage_path"]))
        else:
            segments = segment_document(document["extracted_text"])
        translated_by_id: dict[str, str] = {}
        detected_language: str | None = None
        resolved_source = options["source_language"]
        batches = list(iter_segment_batches(segments))
        if progress_callback:
            progress_callback(0, len(batches))
        logger.info(
            "translation_started request_id=%s document_id=%s segment_count=%d batch_count=%d",
            request_id,
            document_id,
            len(segments),
            len(batches),
        )
        for batch_index, batch in enumerate(batches, 1):
            detected, translated = await translate_batch_with_recovery(
                batch,
                source_language=resolved_source,
                target_language=options["target_language"],
                style=options["style"],
                glossary=options["glossary"],
                user_id=user_id,
                request_id=request_id,
                batch_index=batch_index,
            )
            if detected_language is None and detected:
                detected_language = detected
                if resolved_source == "auto":
                    resolved_source = detected
            translated_by_id.update(
                (item["segment_id"], item["translated_text"]) for item in translated
            )
            if progress_callback:
                progress_callback(batch_index, len(batches))
        result_segments = [
            {**segment, "translated_text": translated_by_id[segment["segment_id"]]}
            for segment in segments
        ]
        logger.info(
            "translation_batches_completed request_id=%s document_id=%s "
            "segment_count=%d batch_count=%d",
            request_id,
            document_id,
            len(result_segments),
            len(batches),
        )
        result = {
            "source_language": options["source_language"],
            "detected_source_language": detected_language,
            "target_language": options["target_language"],
            "mode": options["mode"],
            "style": options["style"],
            "segments": result_segments,
        }
        _finish_translation(
            task["id"],
            user_id,
            result=result,
            detected_source_language=detected_language,
        )
        return {
            "document_id": document_id,
            "translation_id": task["id"],
            **result,
            "cached": False,
        }
    except DocumentTranslationError as exc:
        _finish_translation(task["id"], user_id, error=str(exc))
        raise
    except Exception as exc:
        error = "文档翻译失败，请重试"
        _finish_translation(task["id"], user_id, error=error)
        raise DocumentTranslationError(error) from exc


def _translation_job_row(job_id: str, user_id: str | None = None):
    statement = select(document_translation_jobs).where(
        document_translation_jobs.c.id == job_id
    )
    if user_id is not None:
        statement = statement.where(document_translation_jobs.c.user_id == user_id)
    with _engine.connect() as conn:
        return conn.execute(statement).mappings().first()


def cleanup_translation_jobs(now: int | None = None) -> dict[str, int]:
    """Expire abandoned jobs and delete terminal job status after its TTL."""
    init_user_database()
    current_time = int(time.time()) if now is None else now
    stale_before = current_time - TRANSLATION_JOB_TIMEOUT_SECONDS
    with _engine.begin() as conn:
        expired_running = conn.execute(
            update(document_translation_jobs)
            .where(
                document_translation_jobs.c.status.in_(("pending", "processing")),
                document_translation_jobs.c.updated_at < stale_before,
            )
            .values(
                status="failed",
                error_code="TRANSLATION_JOB_TIMEOUT",
                error_message="文档翻译任务执行超时，请重新提交",
                completed_at=current_time,
                updated_at=current_time,
                expires_at=current_time + TRANSLATION_JOB_TTL_SECONDS,
            )
        ).rowcount
        deleted = conn.execute(
            delete(document_translation_jobs).where(
                document_translation_jobs.c.status.in_(("completed", "failed")),
                document_translation_jobs.c.expires_at <= current_time,
            )
        ).rowcount
    return {"expired": expired_running, "deleted": deleted}


def _job_progress(row: dict[str, Any]) -> dict[str, int]:
    total = max(0, int(row.get("total_batches") or 0))
    completed = min(total, max(0, int(row.get("completed_batches") or 0)))
    percent = (
        100
        if row.get("status") == "completed"
        else (round(completed * 100 / total) if total else 0)
    )
    return {
        "completed_batches": completed,
        "total_batches": total,
        "percent": percent,
    }


def _translation_job_response(row: dict[str, Any]) -> dict[str, Any]:
    response: dict[str, Any] = {
        "job_id": row["id"],
        "document_id": row["document_id"],
        "status": row["status"],
        "progress": _job_progress(row),
        "request_id": row["request_id"],
        "error_code": row.get("error_code"),
        "message": row.get("error_message"),
        "expires_at": row["expires_at"],
    }
    if row["status"] == "completed" and row.get("translation_id"):
        _, result = get_translation(
            row["user_id"], row["document_id"], row["translation_id"]
        )
        response["result"] = {
            "document_id": row["document_id"],
            "translation_id": row["translation_id"],
            **result,
            "cached": bool(row.get("cached")),
        }
    return response


def create_translation_job(
    user_id: str,
    document_id: str,
    *,
    source_language: str = "auto",
    target_language: str,
    mode: str = "translation",
    style: str = "general",
    glossary: Any = None,
) -> tuple[dict[str, Any], bool]:
    """Create or reuse one persistent job for a document/options combination."""
    document = get_owned_document(user_id, document_id)
    if document["extraction_status"] != "completed" or not document["extracted_text"]:
        raise DocumentTranslationError("请先完成文档解析")
    options = normalize_options(
        source_language=source_language,
        target_language=target_language,
        mode=mode,
        style=style,
        glossary=glossary,
    )
    if document.get("file_type") == "pdf" and options["mode"] != "translation":
        raise InvalidTranslationRequestError("PDF翻译暂不支持双语模式")
    if document.get("file_type") == "pdf":
        segments = segment_pdf_document(Path(document["storage_path"]))
    elif document.get("file_type") == "docx":
        segments = segment_docx_document(Path(document["storage_path"]))
    else:
        segments = segment_document(document["extracted_text"])
    total_batches = max(1, len(list(iter_segment_batches(segments))))
    options_hash = _options_hash(options)
    options_json = json.dumps(options, ensure_ascii=False, separators=(",", ":"))
    now = int(time.time())
    cleanup_translation_jobs(now)

    with _engine.connect() as conn:
        existing = (
            conn.execute(
                select(document_translation_jobs).where(
                    document_translation_jobs.c.document_id == document_id,
                    document_translation_jobs.c.user_id == user_id,
                    document_translation_jobs.c.options_hash == options_hash,
                )
            )
            .mappings()
            .first()
        )
    if existing and existing["status"] in ("pending", "processing", "completed"):
        return _translation_job_response(dict(existing)), False

    request_id = uuid.uuid4().hex
    values = {
        "id": existing["id"] if existing else str(uuid.uuid4()),
        "document_id": document_id,
        "user_id": user_id,
        "options_hash": options_hash,
        "options_json": options_json,
        "translation_id": None,
        "request_id": request_id,
        "status": "pending",
        "completed_batches": 0,
        "total_batches": total_batches,
        "cached": 0,
        "error_code": None,
        "error_message": None,
        "created_at": now,
        "started_at": None,
        "updated_at": now,
        "completed_at": None,
        "expires_at": now
        + TRANSLATION_JOB_TIMEOUT_SECONDS
        + TRANSLATION_JOB_TTL_SECONDS,
    }
    try:
        with _engine.begin() as conn:
            if existing:
                conn.execute(
                    update(document_translation_jobs)
                    .where(document_translation_jobs.c.id == existing["id"])
                    .values(
                        **{key: value for key, value in values.items() if key != "id"}
                    )
                )
            else:
                conn.execute(insert(document_translation_jobs).values(**values))
    except IntegrityError:
        with _engine.connect() as conn:
            winner = (
                conn.execute(
                    select(document_translation_jobs).where(
                        document_translation_jobs.c.document_id == document_id,
                        document_translation_jobs.c.user_id == user_id,
                        document_translation_jobs.c.options_hash == options_hash,
                    )
                )
                .mappings()
                .first()
            )
        if not winner:
            raise
        return _translation_job_response(dict(winner)), False
    row = _translation_job_row(values["id"], user_id)
    return _translation_job_response(dict(row)), True


def get_translation_job(user_id: str, job_id: str) -> dict[str, Any]:
    cleanup_translation_jobs()
    row = _translation_job_row(job_id, user_id)
    if not row:
        raise KeyError("翻译任务不存在或已过期")
    return _translation_job_response(dict(row))


def _set_translation_job_progress(job_id: str, completed: int, total: int) -> None:
    now = int(time.time())
    with _engine.begin() as conn:
        conn.execute(
            update(document_translation_jobs)
            .where(
                document_translation_jobs.c.id == job_id,
                document_translation_jobs.c.status == "processing",
            )
            .values(
                completed_batches=max(0, completed),
                total_batches=max(1, total),
                updated_at=now,
            )
        )


def _translation_job_error_code(exc: Exception) -> str:
    message = str(exc)
    if "超时" in message:
        return "DEEPSEEK_TIMEOUT"
    if "网络连接" in message:
        return "DEEPSEEK_CONNECTION_FAILED"
    if "DeepSeek API 请求失败" in message:
        return "DEEPSEEK_HTTP_ERROR"
    if "AI 返回" in message or "翻译失败" in message:
        return "TRANSLATION_RESPONSE_INVALID"
    return "TRANSLATION_FAILED"


def _fail_translation_job(
    job_id: str, *, error_code: str, message: str, request_id: str
) -> None:
    now = int(time.time())
    logger.error(
        "translation_job_failed job_id=%s request_id=%s error_code=%s message=%s",
        job_id,
        request_id,
        error_code,
        message,
    )
    job_row = _translation_job_row(job_id)
    with _engine.begin() as conn:
        conn.execute(
            update(document_translation_jobs)
            .where(document_translation_jobs.c.id == job_id)
            .values(
                status="failed",
                error_code=error_code,
                error_message=message[:1000],
                completed_at=now,
                updated_at=now,
                expires_at=now + TRANSLATION_JOB_TTL_SECONDS,
            )
        )
        if job_row:
            # asyncio timeout / cancellation can interrupt translate_document before
            # it marks its cache row failed. Release that row so a retry can start.
            conn.execute(
                update(document_translations)
                .where(
                    document_translations.c.document_id == job_row["document_id"],
                    document_translations.c.user_id == job_row["user_id"],
                    document_translations.c.options_hash == job_row["options_hash"],
                    document_translations.c.status == "processing",
                )
                .values(
                    status="failed",
                    error_message=message[:1000],
                    updated_at=now,
                )
            )


async def run_translation_job(job_id: str) -> None:
    """Claim and execute one job; every failure is persisted and contained."""
    row = _translation_job_row(job_id)
    if not row:
        return
    now = int(time.time())
    with _engine.begin() as conn:
        claimed = conn.execute(
            update(document_translation_jobs)
            .where(
                document_translation_jobs.c.id == job_id,
                document_translation_jobs.c.status == "pending",
            )
            .values(status="processing", started_at=now, updated_at=now)
        )
    if claimed.rowcount != 1:
        return

    row = dict(_translation_job_row(job_id))
    request_id = row["request_id"]
    try:
        options = json.loads(row["options_json"])
        result = await asyncio.wait_for(
            translate_document(
                row["user_id"],
                row["document_id"],
                source_language=options["source_language"],
                target_language=options["target_language"],
                mode=options["mode"],
                style=options["style"],
                glossary=options["glossary"],
                request_id=request_id,
                progress_callback=lambda completed, total: _set_translation_job_progress(
                    job_id, completed, total
                ),
            ),
            timeout=TRANSLATION_JOB_TIMEOUT_SECONDS,
        )
        finished_at = int(time.time())
        current = dict(_translation_job_row(job_id))
        total_batches = max(1, int(current.get("total_batches") or 1))
        with _engine.begin() as conn:
            conn.execute(
                update(document_translation_jobs)
                .where(document_translation_jobs.c.id == job_id)
                .values(
                    status="completed",
                    translation_id=result["translation_id"],
                    completed_batches=total_batches,
                    total_batches=total_batches,
                    cached=1 if result.get("cached") else 0,
                    error_code=None,
                    error_message=None,
                    completed_at=finished_at,
                    updated_at=finished_at,
                    expires_at=finished_at + TRANSLATION_JOB_TTL_SECONDS,
                )
            )
    except asyncio.TimeoutError:
        _fail_translation_job(
            job_id,
            error_code="TRANSLATION_JOB_TIMEOUT",
            message="文档翻译任务执行超时，请重新提交",
            request_id=request_id,
        )
    except asyncio.CancelledError:
        _fail_translation_job(
            job_id,
            error_code="TRANSLATION_JOB_CANCELLED",
            message="文档翻译任务已中止，请重新提交",
            request_id=request_id,
        )
        raise
    except DocumentTranslationError as exc:
        _fail_translation_job(
            job_id,
            error_code=_translation_job_error_code(exc),
            message=str(exc),
            request_id=request_id,
        )
    except Exception as exc:
        logger.exception(
            "translation_job_unhandled_error job_id=%s request_id=%s",
            job_id,
            request_id,
        )
        _fail_translation_job(
            job_id,
            error_code="TRANSLATION_INTERNAL_ERROR",
            message="文档翻译任务执行失败，请稍后重试",
            request_id=request_id,
        )


async def _translation_job_cleanup_loop() -> None:
    try:
        while True:
            await asyncio.sleep(TRANSLATION_JOB_CLEANUP_INTERVAL_SECONDS)
            try:
                cleanup_translation_jobs()
            except Exception:
                logger.exception("translation_job_cleanup_failed")
    except asyncio.CancelledError:
        return


def _discard_background_job(task: asyncio.Task[None]) -> None:
    _background_jobs.discard(task)
    if task.cancelled():
        return
    try:
        task.result()
    except Exception:
        logger.exception("translation_background_task_escaped")


def schedule_translation_job(job_id: str) -> None:
    global _cleanup_task
    task = asyncio.create_task(run_translation_job(job_id))
    _background_jobs.add(task)
    task.add_done_callback(_discard_background_job)
    if _cleanup_task is None or _cleanup_task.done():
        _cleanup_task = asyncio.create_task(_translation_job_cleanup_loop())


def get_translation(
    user_id: str, document_id: str, translation_id: str
) -> tuple[dict[str, Any], dict[str, Any]]:
    document = get_owned_document(user_id, document_id)
    init_user_database()
    with _engine.connect() as conn:
        row = (
            conn.execute(
                select(document_translations).where(
                    document_translations.c.id == translation_id,
                    document_translations.c.document_id == document_id,
                    document_translations.c.user_id == user_id,
                )
            )
            .mappings()
            .first()
        )
    if not row:
        raise KeyError("翻译结果不存在")
    if row["status"] != "completed" or not row["result_json"]:
        raise DocumentTranslationError("翻译结果尚未生成")
    return dict(document), json.loads(row["result_json"])


def _docx_story_container(document, location: dict[str, Any]):
    section = document.sections[int(location["section_index"])]
    variant = location["variant"]
    story = "header" if location["scope"].startswith("header") else "footer"
    if story == "header":
        return (
            section.first_page_header
            if variant == "first"
            else section.even_page_header if variant == "even" else section.header
        )
    return (
        section.first_page_footer
        if variant == "first"
        else section.even_page_footer if variant == "even" else section.footer
    )


def _resolve_docx_location(document, location: dict[str, Any]):
    try:
        scope = location["scope"]
        if scope == "body":
            return document.paragraphs[int(location["paragraph_index"])]
        if scope == "table":
            tables = document.tables
        else:
            container = _docx_story_container(document, location)
            if scope in ("header", "footer"):
                return container.paragraphs[int(location["paragraph_index"])]
            tables = container.tables
        table = tables[int(location["table_index"])]
        row = table.rows[int(location["row_index"])]
        cell = row.cells[int(location["cell_index"])]
        return cell.paragraphs[int(location["paragraph_index"])]
    except (IndexError, KeyError, TypeError, ValueError) as exc:
        raise DocumentTranslationError(
            f"DOCX结构位置无效：{json.dumps(location, ensure_ascii=False)}"
        ) from exc


def _assign_docx_text_nodes(text_nodes: list[Any], translated_text: str) -> None:
    if not text_nodes:
        raise DocumentTranslationError("DOCX目标段落缺少可回填文本节点")
    original_lengths = [max(1, len(node.text or "")) for node in text_nodes]
    remaining_text = translated_text
    remaining_weight = sum(original_lengths)
    for index, (node, weight) in enumerate(zip(text_nodes, original_lengths)):
        if index == len(text_nodes) - 1:
            value = remaining_text
        else:
            take = round(len(remaining_text) * weight / max(1, remaining_weight))
            value = remaining_text[:take]
            remaining_text = remaining_text[take:]
            remaining_weight -= weight
        node.text = value
        xml_space = "{http://www.w3.org/XML/1998/namespace}space"
        if value.startswith(" ") or value.endswith(" "):
            node.set(xml_space, "preserve")
        elif xml_space in node.attrib:
            del node.attrib[xml_space]


def _replace_docx_paragraph_text(paragraph, translated_text: str) -> None:
    tokens = _docx_paragraph_tokens(paragraph)
    text_groups: list[list[Any]] = [[]]
    for token_type, node in tokens:
        if token_type == "text":
            text_groups[-1].append(node)
        else:
            text_groups.append([])
    non_empty_groups = [group for group in text_groups if group]
    if not non_empty_groups:
        raise DocumentTranslationError("DOCX目标段落缺少可回填文本节点")

    translated_parts = re.split(r"[\t\r\n]", translated_text)
    if len(translated_parts) == len(text_groups) and all(text_groups):
        for group, part in zip(text_groups, translated_parts, strict=False):
            _assign_docx_text_nodes(group, part)
        return

    # The model may omit structural separators. Keep the original w:tab / w:br
    # nodes and distribute only translated characters across existing w:t nodes.
    flattened_nodes = [node for group in text_groups for node in group]
    plain_text = re.sub(r"[\t\r\n]", "", translated_text)
    _assign_docx_text_nodes(flattened_nodes, plain_text)


def _render_docx_from_original(
    document_row: dict[str, Any], segments: list[dict[str, Any]]
) -> bytes:
    source_path = Path(str(document_row.get("storage_path") or ""))
    if document_row.get("file_type") != "docx" or not source_path.is_file():
        raise InvalidTranslationRequestError("只有DOCX原文件可以导出译文DOCX")
    if not segments or any(
        not isinstance(segment.get("location"), dict) for segment in segments
    ):
        raise InvalidTranslationRequestError(
            "旧DOCX翻译结果缺少结构位置，请重新翻译后再导出"
        )

    output_document = Document(str(source_path))
    backfilled = 0
    try:
        for segment in segments:
            paragraph = _resolve_docx_location(output_document, segment["location"])
            _replace_docx_paragraph_text(paragraph, segment["translated_text"])
            backfilled += 1
    except DocumentTranslationError:
        raise
    except Exception as exc:
        logger.exception(
            "translation_docx_backfill_failed document_id=%s segment_id=%s",
            document_row.get("id"),
            segment.get("segment_id"),
        )
        raise DocumentTranslationError(
            f"DOCX原格式回填失败：{segment.get('segment_id')}"
        ) from exc
    if backfilled != len(segments):
        raise DocumentTranslationError(
            f"DOCX原格式回填不完整：{backfilled}/{len(segments)}"
        )

    stream = io.BytesIO()
    try:
        output_document.save(stream)
    except Exception as exc:
        raise DocumentTranslationError("DOCX原格式文件生成失败") from exc
    return stream.getvalue()


def _pdf_segments_with_layout(
    document_row: dict[str, Any], segments: list[dict[str, Any]]
) -> list[dict[str, Any]]:
    if all(isinstance(segment.get("layout"), dict) for segment in segments):
        return segments
    raise InvalidTranslationRequestError(
        "旧PDF翻译结果缺少版式信息，请重新翻译后再导出PDF"
    )


def _apply_pdf_text_redactions(pdf_document, segments: list[dict[str, Any]]) -> None:
    import fitz

    by_page: dict[int, list[dict[str, Any]]] = {}
    for segment in segments:
        page_number = int(segment["layout"]["page_number"])
        by_page.setdefault(page_number, []).append(segment)

    for page_number, page_segments in by_page.items():
        if page_number < 1 or page_number > pdf_document.page_count:
            raise DocumentTranslationError(f"PDF 第 {page_number} 页不存在")
        page = pdf_document[page_number - 1]
        for segment in page_segments:
            rect = fitz.Rect(segment["layout"]["bbox"])
            if rect.is_empty or rect.is_infinite:
                raise DocumentTranslationError(
                    f"PDF 文本块 {segment['segment_id']} 坐标无效"
                )
            page.add_redact_annot(rect, fill=None, cross_out=False)
        redaction_options = {
            "images": getattr(fitz, "PDF_REDACT_IMAGE_NONE", 0),
            "graphics": getattr(fitz, "PDF_REDACT_LINE_ART_NONE", 0),
        }
        try:
            page.apply_redactions(
                **redaction_options,
                text=getattr(fitz, "PDF_REDACT_TEXT_REMOVE", 0),
            )
        except TypeError:  # PyMuPDF versions before the text option was introduced.
            page.apply_redactions(**redaction_options)


def _normalize_pdf_translated_text(text: str, kind: str) -> str:
    """Remove physical PDF line wraps while retaining explicit paragraph breaks."""
    normalized = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    if kind == "heading":
        return re.sub(r"\s+", " ", normalized).strip()
    paragraphs = re.split(r"\n[ \t]*\n+", normalized)
    normalized_paragraphs = [
        re.sub(r"\s+", " ", paragraph).strip() for paragraph in paragraphs
    ]
    return "\n\n".join(paragraph for paragraph in normalized_paragraphs if paragraph)


def _pdf_translation_html(normalized_text: str) -> str:
    paragraphs = [html.escape(part) for part in normalized_text.split("\n\n")]
    return f"<div>{'<br><br>'.join(paragraphs)}</div>"


def _pdf_translation_css(layout: dict[str, Any]) -> str:
    font_size = max(4.0, float(layout.get("font_size") or 10.0))
    font_color = str(layout.get("font_color") or "#000000")
    font_family = str(layout.get("font_family") or "sans-serif")
    text_align = str(layout.get("text_align") or "left")
    return (
        "* { margin: 0; padding: 0; } "
        f"body {{ font-family: {font_family}; font-size: {font_size:.2f}pt; "
        f"line-height: 1.08; color: {font_color}; text-align: {text_align}; }}"
    )


def _pdf_rect_values(rect) -> list[float]:
    return [round(float(value), 3) for value in (rect.x0, rect.y0, rect.x1, rect.y1)]


def _pdf_inflate_rect(rect, amount: float):
    import fitz

    return fitz.Rect(
        rect.x0 - amount,
        rect.y0 - amount,
        rect.x1 + amount,
        rect.y1 + amount,
    )


def _pdf_drawing_obstacles(page) -> list[Any]:
    """Return thin geometry obstacles instead of blocking a drawing's whole bounds."""
    import fitz

    obstacles: list[Any] = []
    for drawing in page.get_drawings():
        stroke = max(0.5, float(drawing.get("width") or 1.0) / 2)
        for item in drawing.get("items") or []:
            item_type = item[0]
            if item_type == "re" and len(item) > 1:
                rect = fitz.Rect(item[1])
                obstacles.extend(
                    [
                        fitz.Rect(
                            rect.x0 - stroke,
                            rect.y0 - stroke,
                            rect.x1 + stroke,
                            rect.y0 + stroke,
                        ),
                        fitz.Rect(
                            rect.x0 - stroke,
                            rect.y1 - stroke,
                            rect.x1 + stroke,
                            rect.y1 + stroke,
                        ),
                        fitz.Rect(rect.x0 - stroke, rect.y0, rect.x0 + stroke, rect.y1),
                        fitz.Rect(rect.x1 - stroke, rect.y0, rect.x1 + stroke, rect.y1),
                    ]
                )
                continue
            points = [
                value
                for value in item[1:]
                if hasattr(value, "x") and hasattr(value, "y")
            ]
            if points:
                x_values = [float(point.x) for point in points]
                y_values = [float(point.y) for point in points]
                obstacles.append(
                    fitz.Rect(
                        min(x_values) - stroke,
                        min(y_values) - stroke,
                        max(x_values) + stroke,
                        max(y_values) + stroke,
                    )
                )
    return obstacles


def _pdf_page_obstacles(
    page, page_segments: list[dict[str, Any]], current_segment_id: str
) -> list[Any]:
    import fitz

    obstacles = [
        fitz.Rect(segment["layout"]["bbox"])
        for segment in page_segments
        if segment["segment_id"] != current_segment_id
    ]
    for image in page.get_image_info(xrefs=True):
        bbox = image.get("bbox")
        if bbox:
            obstacles.append(fitz.Rect(bbox))
    obstacles.extend(_pdf_drawing_obstacles(page))
    return [rect for rect in obstacles if not rect.is_empty and not rect.is_infinite]


def _pdf_rect_has_new_collision(original, candidate, obstacle) -> bool:
    padded = _pdf_inflate_rect(obstacle, PDF_LAYOUT_GAP)
    return candidate.intersects(padded) and not original.intersects(padded)


def _expand_pdf_bbox_down(page_rect, original, obstacles):
    import fitz

    bottom = float(page_rect.y1)
    for obstacle in obstacles:
        padded = _pdf_inflate_rect(obstacle, PDF_LAYOUT_GAP)
        horizontally_overlaps = not (
            padded.x1 <= original.x0 or padded.x0 >= original.x1
        )
        if horizontally_overlaps and padded.y0 >= original.y1:
            bottom = min(bottom, float(padded.y0))
    bottom = max(float(original.y1), bottom)
    return fitz.Rect(original.x0, original.y0, original.x1, bottom)


def _expand_pdf_bbox_horizontally(page_rect, original, vertically_expanded, obstacles):
    import fitz

    expansion = min(
        PDF_LAYOUT_HORIZONTAL_EXPANSION,
        max(6.0, float(original.width) * 0.12),
    )
    left = max(float(page_rect.x0), float(original.x0) - expansion)
    right = min(float(page_rect.x1), float(original.x1) + expansion)
    candidate = fitz.Rect(
        left,
        original.y0,
        right,
        vertically_expanded.y1,
    )
    for obstacle in obstacles:
        if not _pdf_rect_has_new_collision(original, candidate, obstacle):
            continue
        padded = _pdf_inflate_rect(obstacle, PDF_LAYOUT_GAP)
        if padded.x1 <= original.x0:
            left = max(left, float(padded.x1))
        elif padded.x0 >= original.x1:
            right = min(right, float(padded.x0))
        else:
            return vertically_expanded
        candidate = fitz.Rect(left, original.y0, right, vertically_expanded.y1)
    if right <= left:
        return vertically_expanded
    return candidate


def _pdf_layout_attempts(page, segment: dict[str, Any], page_segments):
    import fitz

    original = fitz.Rect(segment["layout"]["bbox"])
    obstacles = _pdf_page_obstacles(page, page_segments, segment["segment_id"])
    downward = _expand_pdf_bbox_down(page.rect, original, obstacles)
    horizontal = _expand_pdf_bbox_horizontally(page.rect, original, downward, obstacles)
    attempts = [
        ("original", original, 1.0),
        ("scaled", original, PDF_TRANSLATION_MIN_SCALE),
    ]
    if downward.y1 > original.y1 + 0.01:
        attempts.append(("expanded_down", downward, PDF_TRANSLATION_MIN_SCALE))
    if horizontal.x0 < downward.x0 - 0.01 or horizontal.x1 > downward.x1 + 0.01:
        attempts.append(("expanded_horizontal", horizontal, PDF_TRANSLATION_MIN_SCALE))
    final_rect = horizontal if horizontal != original else downward
    attempts.append(
        ("absolute_minimum", final_rect, PDF_TRANSLATION_ABSOLUTE_MIN_SCALE)
    )

    unique_attempts = []
    seen: set[tuple[float, ...]] = set()
    for stage, rect, scale_low in attempts:
        key = (*[round(value, 3) for value in rect], round(scale_low, 3))
        if key in seen:
            continue
        seen.add(key)
        unique_attempts.append((stage, rect, scale_low))
    return unique_attempts


def _fit_pdf_layout_block(page, segment, page_segments) -> dict[str, Any]:
    layout = segment["layout"]
    original_bbox = list(layout["bbox"])
    normalized_text = _normalize_pdf_translated_text(
        segment["translated_text"], segment.get("kind", "paragraph")
    )
    content = _pdf_translation_html(normalized_text)
    css = _pdf_translation_css(layout)
    last_rect = None
    last_scale = PDF_TRANSLATION_ABSOLUTE_MIN_SCALE
    for stage, rect, scale_low in _pdf_layout_attempts(page, segment, page_segments):
        last_rect = rect
        last_scale = scale_low
        if stage.startswith("expanded"):
            logger.info(
                "pdf_layout_bbox_expanded page_number=%s segment_id=%s stage=%s "
                "original_bbox=%s attempted_bbox=%s",
                layout["page_number"],
                segment["segment_id"],
                stage,
                original_bbox,
                _pdf_rect_values(rect),
            )
        spare_height, fitted_scale = page.insert_htmlbox(
            rect,
            content,
            css=css,
            scale_low=scale_low,
            overlay=True,
        )
        if spare_height >= 0:
            logger.info(
                "pdf_layout_block_fitted page_number=%s segment_id=%s stage=%s "
                "bbox=%s scale=%.3f",
                layout["page_number"],
                segment["segment_id"],
                stage,
                _pdf_rect_values(rect),
                fitted_scale,
            )
            return {
                "segment": segment,
                "bbox": _pdf_rect_values(rect),
                "scale_low": scale_low,
                "fitted_scale": fitted_scale,
                "normalized_text": normalized_text,
                "css": css,
            }

    attempted_bbox = _pdf_rect_values(last_rect)
    error = PDFLayoutOverflowError(
        page_number=int(layout["page_number"]),
        segment_id=segment["segment_id"],
        original_bbox=original_bbox,
        attempted_bbox=attempted_bbox,
        attempted_scale=last_scale,
        translated_length=len(str(segment["translated_text"])),
    )
    logger.warning(
        "pdf_layout_overflow error_code=%s page_number=%s segment_id=%s "
        "original_bbox=%s attempted_bbox=%s attempted_scale=%.3f translated_length=%s",
        error.error_code,
        error.page_number,
        error.segment_id,
        error.original_bbox,
        error.attempted_bbox,
        error.attempted_scale,
        error.translated_length,
    )
    raise error


def _preflight_pdf_layout(pdf_document, segments: list[dict[str, Any]]):
    by_page: dict[int, list[dict[str, Any]]] = {}
    for segment in segments:
        by_page.setdefault(int(segment["layout"]["page_number"]), []).append(segment)
    logger.info("pdf_layout_preflight_started segment_count=%s", len(segments))
    plans = []
    for segment in segments:
        page_number = int(segment["layout"]["page_number"])
        if page_number < 1 or page_number > pdf_document.page_count:
            raise DocumentTranslationError(f"PDF 第 {page_number} 页不存在")
        plans.append(
            _fit_pdf_layout_block(
                pdf_document[page_number - 1], segment, by_page[page_number]
            )
        )
    return plans


def _insert_pdf_translations(pdf_document, plans: list[dict[str, Any]]) -> None:
    import fitz

    for plan in plans:
        segment = plan["segment"]
        layout = segment["layout"]
        page = pdf_document[int(layout["page_number"]) - 1]
        rect = fitz.Rect(plan["bbox"])
        spare_height, _ = page.insert_htmlbox(
            rect,
            _pdf_translation_html(plan["normalized_text"]),
            css=plan["css"],
            scale_low=plan["scale_low"],
            overlay=True,
        )
        if spare_height < 0:
            raise PDFLayoutOverflowError(
                page_number=int(layout["page_number"]),
                segment_id=segment["segment_id"],
                original_bbox=layout["bbox"],
                attempted_bbox=plan["bbox"],
                attempted_scale=plan["scale_low"],
                translated_length=len(str(segment["translated_text"])),
            )


def _render_pdf_from_original(
    document_row: dict[str, Any], result: dict[str, Any]
) -> bytes:
    if document_row.get("file_type") != "pdf":
        raise InvalidTranslationRequestError("只有PDF原文件可以导出原版式PDF")
    if result.get("mode") != "translation":
        raise InvalidTranslationRequestError("PDF翻译暂不支持双语模式")
    source_path = Path(str(document_row.get("storage_path") or ""))
    if not source_path.is_file():
        raise DocumentTranslationError("PDF 原文件不存在")

    import fitz

    segments = _pdf_segments_with_layout(document_row, result["segments"])
    try:
        source_bytes = source_path.read_bytes()
        with fitz.open(stream=source_bytes, filetype="pdf") as preflight_document:
            if preflight_document.needs_pass:
                raise DocumentTranslationError("暂不支持加密 PDF 的原版式翻译")
            plans = _preflight_pdf_layout(preflight_document, segments)
        with fitz.open(stream=source_bytes, filetype="pdf") as output_document:
            _apply_pdf_text_redactions(output_document, segments)
            _insert_pdf_translations(output_document, plans)
            return output_document.tobytes(garbage=4, deflate=True)
    except DocumentTranslationError:
        raise
    except Exception as exc:
        logger.exception(
            "translation_pdf_export_failed document_id=%s translation_mode=%s",
            document_row.get("id"),
            result.get("mode"),
        )
        raise DocumentTranslationError("原版式 PDF 生成失败") from exc


def render_translation_export(
    user_id: str, document_id: str, translation_id: str, export_format: str
) -> tuple[bytes, str, str]:
    document, result = get_translation(user_id, document_id, translation_id)
    file_type = document.get("file_type")
    if file_type == "pdf":
        if export_format != "pdf":
            raise InvalidTranslationRequestError("PDF翻译结果仅支持导出为PDF")
        payload = _render_pdf_from_original(document, result)
        media_type = "application/pdf"
    elif file_type == "docx":
        if export_format != "docx":
            raise InvalidTranslationRequestError("DOCX翻译结果仅支持导出为DOCX")
        payload = _render_docx_from_original(document, result["segments"])
        media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        raise InvalidTranslationRequestError("该文档格式不支持翻译导出")
    base_name = document["filename"].rsplit(".", 1)[0] or "translation"
    return payload, media_type, f"{base_name}-translated.{export_format}"
